import os
from huggingface_hub import snapshot_download
import json
import re
import glob
import numpy as np
import faiss
import markdown
from datetime import datetime
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
from sentence_transformers import SentenceTransformer
import PyPDF2
import streamlit as st
from PIL import Image
import streamlit.components.v1 as components
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from typing import List, Dict, Any, Tuple, Optional
import matplotlib.pyplot as plt
from nltk.sentiment.vader import SentimentIntensityAnalyzer
import pandas as pd
from io import BytesIO
import base64
import requests
import plotly.express as px
from collections import Counter
import unicodedata # Import unicodedata for advanced cleaning


from huggingface_hub import snapshot_download
snapshot_download(repo_id="sentence-transformers/all-MiniLM-L6-v2")

MODEL_DIR = "local_models/sentence-transformers/all-MiniLM-L6-v2"
# Helper function to remove problematic Unicode characters
def remove_problematic_chars(text):
    """Removes characters that might cause encoding or display issues,
       especially those outside common encodings like latin-1, by replacing
       common problematic characters and filtering others."""
    if not isinstance(text, str):
        return text # Return as is if not a string

    # Replace common problematic characters with safe alternatives
    cleaned_text = text.replace('\u2013', '-') # En dash
    cleaned_text = cleaned_text.replace('\u2014', '-') # Em dash
    cleaned_text = cleaned_text.replace('\u2018', "'") # Left single quote
    cleaned_text = cleaned_text.replace('\u2019', "'") # Right single quote (apostrophe)
    cleaned_text = cleaned_text.replace('\u201c', '"') # Left double quote
    cleaned_text = cleaned_text.replace('\u201d', '"') # Right double quote
    cleaned_text = cleaned_text.replace('\u2026', '...') # Ellipsis
    cleaned_text = cleaned_text.replace('\u2022', '*') # Bullet point
    cleaned_text = cleaned_text.replace('\u2122', '(TM)') # Trade Mark symbol
    cleaned_text = cleaned_text.replace('\u00AE', '(R)') # Registered symbol
    cleaned_text = cleaned_text.replace('\u00A9', '(C)') # Copyright symbol


    # Attempt standard encode/decode with 'ignore' errors for characters still problematic
    # This is a fallback and might lose some characters not explicitly handled above
    try:
        # Try encoding to latin-1 and then decoding. Characters not in latin-1 will be ignored.
        cleaned_text = cleaned_text.encode('latin-1', errors='ignore').decode('latin-1')
    except UnicodeEncodeError:
        # If latin-1 encoding still fails unexpectedly, try a more aggressive UTF-8 encode/decode with replace
        cleaned_text = cleaned_text.encode('utf-8', errors='replace').decode('utf-8')
    except Exception as e:
        # Catch any other potential encoding errors during this step
        print(f"Warning: Error during initial encoding cleanup: {e}. Proceeding with filtering.")


    # Further filter potentially problematic Unicode characters using a stricter regex
    # This regex aims to keep printable ASCII, common whitespace, and a limited set of safe Unicode.
    # It removes control characters, surrogates, and many other non-ASCII characters.
    # Keeping only printable ASCII and basic whitespace:
    # \x20-\x7E : Printable ASCII characters (space to tilde)
    # \n\r\t : Newline, carriage return, tab
    # \u00A0 : Non-breaking space (often useful)
    # \u20AC : Euro sign (common) - can add others if needed
    # This regex is stricter than the previous one to avoid latin-1 issues.
    # Adjust the regex based on what characters are expected/allowed in the output documents.
    # Current regex keeps printable ASCII, newline, carriage return, tab, and non-breaking space.
    problematic_chars_regex = re.compile(r'[^\x20-\x7E\n\r\t\u00A0]')


    filtered_text = problematic_chars_regex.sub('', cleaned_text)

    return filtered_text


# Load configuration
def load_config():
    """Load configuration from config.json or create default if not exists"""
    config_path = "config.json"

    default_config = {
        "company_info": {
            "name": "Your Company Name",
            "logo_path": "",
            "default_styles": {
                "primary_color": "#003366",
                "secondary_color": "#669933",
                "font_family": "Arial"
            }
        },
        "api_keys": {
            "openai_key": os.environ.get("OPENAI_API_KEY", ""),
            "google_api_key": os.environ.get("GOOGLE_API_KEY", "")
        },
        "knowledge_base": {
            "directory": "markdown_responses",
            "embedding_model": "all-MiniLM-L6-v2",
            "metadata_fields": ["client_industry", "proposal_success", "project_size", "key_differentiators"]
        },
        "proposal_settings": {
            "default_sections": [],
            "max_tokens_per_section": 2000,
            "templates": ["Standard RFP", "Technical RFP", "Commercial RFP"]
        },
        "internal_capabilities": {
            "technical": ["Cloud solutions", "AI implementation", "Data analytics"],
            "functional": ["Project management", "24/7 support", "Custom development"]
        },
        "scoring_system": {
            "weighting": {
                "requirement_match": 0.4,
                "compliance": 0.25,
                "quality": 0.2,
                "alignment": 0.15,
                "risk": 0.1
            },
            "grading_scale": {
                "excellent": [90, 100],
                "good": [70, 89],
                "fair": [50, 69],
                "poor": [0, 49]
            }
        }
    }

    if not os.path.exists(config_path):
        print("config.json not found, creating default.")
        with open(config_path, 'w') as f:
            json.dump(default_config, f, indent=4)
        return default_config

    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
            # Merge with default config to ensure all keys exist
            # This handles cases where config.json exists but is missing sections
            merged_config = default_config.copy()
            merged_config.update(config)
            return merged_config
    except json.JSONDecodeError:
        print(f"Error decoding JSON from {config_path}. Using default config.")
        return default_config
    except Exception as e:
        print(f"An unexpected error occurred loading {config_path}: {e}. Using default config.")
        return default_config


# Document processing functions
def extract_text_from_docx(file_path):
    """Extract text from DOCX files including tables and headers"""
    doc = Document(file_path)
    full_text = []

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Apply cleaning to cell text
                cleaned_cell_text = remove_problematic_chars(cell.text.strip())
                if cleaned_cell_text:
                    row_text.append(cleaned_cell_text)
            if row_text:
                full_text.append(" | ".join(row_text))

    for para in doc.paragraphs:
        # Apply cleaning to paragraph text
        cleaned_para_text = remove_problematic_chars(para.text.strip())
        if cleaned_para_text:
            if para.style.name.startswith('Heading'):
                heading_level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                prefix = '#' * heading_level + ' '
                full_text.append(f"{prefix}{cleaned_para_text}")
            else:
                full_text.append(cleaned_para_text)

    return '\n'.join(full_text) # Text is already cleaned


def extract_text_from_pdf(file_path):
    """Extract text from PDF documents"""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = []
        for page in reader.pages:
            # Apply cleaning to extracted page text
            cleaned_page_text = remove_problematic_chars(page.extract_text())
            text.append(cleaned_page_text)
    return '\n'.join(text) # Text is already cleaned


def extract_sections_from_rfp(rfp_text):
    """Extract structured sections from the RFP text with improved pattern matching"""
    # Ensure the input text is cleaned before processing
    cleaned_rfp_text = remove_problematic_chars(rfp_text)

    section_patterns = [
        r'^(?:\d+\.)?(?:\d+\.)?(?:\d+\.)?\s*([A-Z][A-Za-z\s]+)$',
        r'^([A-Z][A-Z\s]+)(?:\:|\.)?\s*$',
        r'^(?:Section|SECTION)\s+\d+\s*[\:\-\.]\s*([A-Za-z\s]+)$'
    ]

    sections = {}
    current_section = "Overview"
    current_content = []

    for line in cleaned_rfp_text.split('\n'):
        matched = False
        for pattern in section_patterns:
            match = re.match(pattern, line.strip())
            if match:
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                    current_content = []

                current_section = match.group(1).strip()
                matched = True
                break

        if not matched:
            current_content.append(line)

    if current_content:
        sections[current_section] = '\n'.join(current_content)

    return sections

def process_rfp(file_path):
    """Extract text from uploaded RFP document"""
    if file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.md') or file_path.endswith('.txt'):
        # Added errors='replace' to handle problematic characters during reading
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            content = file.read()
        return remove_problematic_chars(content) # Clean content after reading
    else:
        raise ValueError("Unsupported file format. Please use DOCX, PDF, TXT or MD file.")

def expand_query(query: str) -> str:
    """Expand query with relevant synonyms and domain-specific terms"""
    domain_specific_terms = {
        "proposal": ["offer", "bid", "solution"],
        "requirements": ["needs", "specifications", "criteria"],
        "implementation": ["deployment", "execution", "rollout"],
        "support": ["maintenance", "service", "assistance"]
    }

    words = query.split()
    expanded_words = []
    for word in words:
        expanded_words.append(word)
        for key, values in domain_specific_terms.items():
            if word.lower() == key:
                expanded_words.extend(values)
            elif word.lower() in values:
                expanded_words.append(key)

    return ' '.join(expanded_words)

class HierarchicalEmbeddingModel:
    """Model for hierarchical embeddings (document and section level)"""
    def __init__(self, model_name: str):
        # --- CHANGE HERE ---
        # Explicitly set the device. Use 'cuda' if you have a configured GPU,
        # otherwise 'cpu' is safer.
        try:
            # Try loading directly to CPU first, often resolves this.
            self.model = SentenceTransformer(model_name, device='cpu')
            print(f"SentenceTransformer loaded on CPU for model: {model_name}")
        except Exception as e:
            print(f"Error loading SentenceTransformer on CPU: {e}. Trying default loading.")
            # Fallback to default loading if CPU fails (might retry meta issue)
            self.model = SentenceTransformer(model_name)

    def encode(self, texts: List[str], level: str = 'section') -> np.ndarray:
        """Generate embeddings with different pooling strategies based on level"""
        # Ensure texts are cleaned before encoding
        cleaned_texts = [remove_problematic_chars(text) for text in texts]

        if level == 'document':
            embeddings = self.model.encode(cleaned_texts, convert_to_tensor=True)
            # Use weighted pooling for document-level embeddings
            weights = np.linspace(0.1, 1.0, len(embeddings))
            weighted_embeddings = embeddings * weights[:, np.newaxis]
            return np.mean(weighted_embeddings, axis=0)
        else:
            return self.model.encode(cleaned_texts)

class ProposalKnowledgeBase:
    def __init__(self, kb_directory="markdown_responses", embedding_model="all-MiniLM-L6-v2"):
        self.kb_directory = kb_directory
        self.embedding_model = embedding_model
        self.model = None  # Initialize model as None
        self.documents = []
        self.section_map = {}
        self.index = None
        self.tfidf_vectorizer = TfidfVectorizer()

        if not os.path.exists(kb_directory):
            os.makedirs(kb_directory)

        self.load_documents()
        self._build_index()

    def load_model(self):
        """Load the SentenceTransformer model with proper error handling"""
        try:
            # Try loading the model with automatic device detection
            self.model = SentenceTransformer(self.embedding_model)
            print(f"Successfully loaded model: {self.embedding_model}")
        except Exception as e:
            print(f"Error loading model: {e}")
            # Fallback to CPU loading
            print("Trying to load model on CPU...")
            try:
                self.model = SentenceTransformer(self.embedding_model, device='cpu')
                print(f"Successfully loaded model on CPU: {self.embedding_model}")
            except Exception as e:
                print(f"Failed to load model even on CPU: {e}")
                # Provide instructions for offline mode
                print("Please ensure you have the model files downloaded locally or check your internet connection.")
                print("For offline use, download the model files from Hugging Face and set the TRANSFORMERS_CACHE environment variable.")
                self.model = None  # Mark model as unavailable

    def load_documents(self):
        """Load all documents from the knowledge base directory"""
        self.documents = []
        self.section_map = {}
        self.metadata = []

        if not os.path.exists(self.kb_directory):
            return

        for filename in os.listdir(self.kb_directory):
            if filename.endswith('.md') or filename.endswith('.txt'):
                file_path = os.path.join(self.kb_directory, filename)
                with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                    content = file.read()

                sections = self._split_into_sections(content)

                for section_name, section_content in sections.items():
                    doc_id = len(self.documents)
                    metadata = {
                        "client_industry": "general",
                        "proposal_success": True,
                        "project_size": "medium",
                        "key_differentiators": ["quality", "experience"]
                    }

                    # Extract metadata from filename
                    if "_success_" in filename:
                        metadata["proposal_success"] = filename.split("_success_")[1].split("_")[0] == "True"
                    if "_industry_" in filename:
                        metadata["client_industry"] = filename.split("_industry_")[1].split("_")[0]
                    if "_size_" in filename:
                        metadata["project_size"] = filename.split("_size_")[1].split("_")[0]

                    self.documents.append({
                        "id": doc_id,
                        "filename": filename,
                        "section_name": section_name,
                        "content": section_content,
                        "metadata": metadata
                    })

                    if section_name not in self.section_map:
                        self.section_map[section_name] = []
                    self.section_map[section_name].append(doc_id)
                    self.metadata.append(metadata)

        # Load the model after documents are loaded
        self.load_model()

        if self.model is None:
            print("Model could not be loaded. Knowledge Base functionality will be limited.")

    # Moved this function inside the class
    def _split_into_sections(self, content):
        """Split a document into sections based on headers"""
        # Input content is assumed to be already cleaned
        sections = {}
        current_section = "Introduction"
        current_content = []

        for line in content.split('\n'):
            if line.startswith('# '):
                if current_content:
                    sections[remove_problematic_chars(current_section)] = '\n'.join(current_content)
                    current_content = []
                current_section = line[2:].strip()
            elif line.startswith('## '):
                if current_content:
                    sections[remove_problematic_chars(current_section)] = '\n'.join(current_content)
                    current_content = []
                current_section = line[3:].strip()
            else:
                current_content.append(line)

        if current_content:
            sections[remove_problematic_chars(current_section)] = '\n'.join(current_content)

        return sections

    def _build_index(self):
        """Build a FAISS index for fast similarity search"""
        if not self.documents:
            return
        # Ensure texts for indexing are cleaned
        texts = [remove_problematic_chars(doc["content"]) for doc in self.documents]
        embeddings = self.model.encode(texts)
        dimension = embeddings.shape[1]
        self.index = faiss.IndexFlatL2(dimension)
        self.index.add(np.array(embeddings).astype('float32'))
        self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(texts)

    def hybrid_search(self, query, k=5):
        """Hybrid search combining dense and sparse retrieval"""
        if not self.index or not self.documents:
            return []
        # Clean the query before encoding and vectorizing
        cleaned_query = remove_problematic_chars(query)
        query_embedding = self.model.encode([cleaned_query])
        dense_scores, dense_indices = self.index.search(np.array(query_embedding).astype('float32'), k)
        query_tfidf = self.tfidf_vectorizer.transform([cleaned_query])
        sparse_scores = cosine_similarity(query_tfidf, self.tfidf_matrix).flatten()
        sparse_indices = np.argsort(-sparse_scores)[:k]
        combined = []
        seen = set()
        for idx in dense_indices[0]:
            if idx not in seen:
                combined.append((dense_scores[0][list(dense_indices[0]).index(idx)], idx))
                seen.add(idx)
        for idx in sparse_indices:
            if idx not in seen:
                combined.append((sparse_scores[idx], idx))
                seen.add(idx)
        combined.sort(key=lambda x: x[0], reverse=True)
        # Ensure document content in results is cleaned
        results = [{"score": float(score), "document": {
            "id": self.documents[idx]["id"],
            "filename": self.documents[idx]["filename"], # Already cleaned
            "section_name": self.documents[idx]["section_name"], # Already cleaned
            "content": remove_problematic_chars(self.documents[idx]["content"]), # Ensure content is cleaned
            "metadata": self.documents[idx]["metadata"] # Metadata should also be cleaned on load
        }} for score, idx in combined[:k]]
        return results

    def get_common_section_names(self, top_n=15):
        return []

    def multi_hop_search(self, initial_query, k=5):
        """Multi-hop retrieval process"""
        if not self.index or not self.documents or not self.model:
            print("Warning: Knowledge Base is not properly initialized for search.")
            return []
        # Clean the initial query
        cleaned_initial_query = remove_problematic_chars(initial_query)
        first = self.hybrid_search(cleaned_initial_query, k=3*k)
        # Ensure content used for refined query is cleaned
        refined_query = cleaned_initial_query + " " + " ".join([remove_problematic_chars(r["document"]["content"])[ :200] for r in first[:3]])
        second = self.hybrid_search(refined_query, k=k)
        all_r = {r["document"]["id"]: r for r in first+second}
        topk = sorted(all_r.values(), key=lambda x: x["score"], reverse=True)[:k]
        return topk

    def get_section_documents(self, section_name):
        # Ensure section name is cleaned for lookup
        cleaned_section_name = remove_problematic_chars(section_name)
        # Ensure returned document content is cleaned
        return [{
            "id": self.documents[idx]["id"],
            "filename": self.documents[idx]["filename"],
            "section_name": self.documents[idx]["section_name"],
            "content": remove_problematic_chars(self.documents[idx]["content"]),
            "metadata": self.documents[idx]["metadata"]
        } for idx in self.section_map.get(cleaned_section_name, [])]

    def get_all_section_names(self):
        # Return cleaned section names
        return [remove_problematic_chars(name) for name in self.section_map.keys()]

    def extract_pricing_from_kb(self) -> List[int]:
        prices = []
        pattern = re.compile(r'(?:₹|Rs\.?)[\s]*([0-9,]+)')
        md_paths = glob.glob(os.path.join(self.kb_directory, '*.md'))
        for path in md_paths:
            # Added errors='replace' here too for reading
            text = open(path, 'r', encoding='utf-8', errors='replace').read()
            text = remove_problematic_chars(text) # Clean text after reading
            # Ensure regex matching is done on cleaned text
            parts = re.split(r'^#{1,3}\s*COMMERCIAL PROPOSAL\s*$', text, flags=re.IGNORECASE | re.MULTILINE)
            if len(parts) < 2:
                continue
            body = re.split(r'^#{1,3}\s+\w', parts[1], flags=re.MULTILINE)[0]
            for m in pattern.finditer(body):
                val = int(m.group(1).replace(',', ''))
                prices.append(val)
        return prices

class SpecialistRAGDrafter:
    def __init__(self, openai_key=None):
        self.client = OpenAI(api_key=openai_key or os.environ.get("OPENAI_API_KEY"))

    def generate_draft(self, section_name, rfp_section_content, relevant_kb_content, client_name):
        # Ensure all input text is cleaned before sending to LLM
        cleaned_section_name = remove_problematic_chars(section_name)
        cleaned_rfp_section_content = remove_problematic_chars(rfp_section_content) if rfp_section_content else ""
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""

        kb_blob = "\n\n".join([
            f"--- {('Very Relevant' if item['score']>0.7 else 'Relevant')} PAST PROPOSAL ---\n"
            f"From: {remove_problematic_chars(item['document']['filename'])} | Section: {remove_problematic_chars(item['document']['section_name'])}\n"
            f"{remove_problematic_chars(item['document']['content'])}"
            for item in relevant_kb_content
        ])

        summary_resp = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role":"system","content":"You’re an expert at summarizing past proposals."},
                {"role":"user","content":
                    f"Summarize the following past-proposal content into 5–7 bullets, focusing on actionable points:\n\n{kb_blob}"
                }
            ],
            temperature=0.0
        )
        # Clean the summarized KB content from the LLM
        summarized_kb = remove_problematic_chars(summary_resp.choices[0].message.content)

        prompt = f"""
        # DRAFT GENERATION FOR {cleaned_section_name}

        ## SECTION CONTENT TO ADDRESS
        {cleaned_rfp_section_content}

        ## SUMMARY OF PAST PROPOSALS
        {summarized_kb}
        """
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2
            )
            # Clean the generated draft text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            return f"Error generating draft for {cleaned_section_name}: {str(e)}"

    def generate_rfp_template(self, company_objectives, template_type):
        # Ensure input text is cleaned before sending to LLM
        cleaned_company_objectives = remove_problematic_chars(company_objectives)
        cleaned_template_type = remove_problematic_chars(template_type)

        prompt = f"""
        # RFP TEMPLATE GENERATION

        Create a comprehensive RFP template based on the following company objectives and template type:

        COMPANY OBJECTIVES:
        {cleaned_company_objectives}

        TEMPLATE TYPE:
        {cleaned_template_type}

        The template should include:
        1. All standard sections for the selected template type
        2. Custom sections based on the company objectives
        3. Clear structure with appropriate headings and subheadings
        4. Placeholder content where applicable
        5. Evaluation criteria section tailored to the objectives
        6. Submission guidelines

        Format as a professional RFP document.
        """
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            # Clean the generated template text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            return f"Error generating RFP template: {str(e)}"

class EnhancedProposalGenerator:
    def __init__(self, knowledge_base, openai_key=None):
        self.kb = knowledge_base
        self.client = OpenAI(api_key=openai_key or os.environ.get("OPENAI_API_KEY"))
        self.rfp_text = None  # Store RFP text for regeneration
        self.drafter = SpecialistRAGDrafter(openai_key)  # Specialist drafter

        # Check if the knowledge base model is available
        if not self.kb or not hasattr(self.kb, 'model') or self.kb.model is None:
            print("Warning: Knowledge Base model is not initialized. Some features may not work.")
    

    def analyze_rfp(self, rfp_text):
        """Comprehensive RFP analysis using the new prompt"""
        # Clean RFP text before storing and sending to LLM
        cleaned_rfp_text = remove_problematic_chars(rfp_text)
        self.rfp_text = cleaned_rfp_text

        prompt = f"""
        You are an expert proposal analyst. Your task is to analyze the following Request for Proposal (RFP) text and extract key information.
        I need a comprehensive, structured analysis of the following Request for Proposal (RFP). Please organize your analysis into the following specific categories with clear headings:

        1. KEY REQUIREMENTS: Extract specific functional and technical requirements that must be addressed, using exact language from the RFP where possible.

        2. DELIVERABLES: List all concrete deliverables explicitly requested in the RFP.

        3. REQUIRED SECTIONS: Identify EXACTLY what sections must be included in the proposal response. Include both main sections and any specified subsections. Use the exact section names from the RFP.

        4. TIMELINE: Extract all dates, deadlines, and milestones mentioned in the RFP.

        5. BUDGET CONSTRAINTS: Note any explicit budget limitations, pricing structures, or financial parameters mentioned.

        6. EVALUATION CRITERIA: Detail how the proposal will be scored or evaluated, including any weighted criteria.

        7. CLIENT PAIN POINTS: Identify specific problems or challenges the client is trying to solve, both explicit and implied.

        8. UNIQUE CONSIDERATIONS: Flag any special requirements, unusual constraints, or differentiating factors that stand out.

        Format your response as a structured analysis with clear headings for each category. Use bullet points for clarity. Extract specific, actionable information rather than general observations.

        RFP TEXT:
        {cleaned_rfp_text}
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2
            )

            # Clean the generated analysis text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error analyzing RFP: {str(e)}")
            return f"Error analyzing RFP: {str(e)}"

    def extract_mandatory_criteria(self, rfp_analysis):
        """Extract mandatory criteria from RFP analysis"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        try:
            requirements_start = cleaned_rfp_analysis.find("KEY REQUIREMENTS") + len("KEY REQUIREMENTS")
            requirements_end = cleaned_rfp_analysis.find("DELIVERABLES", requirements_start)
            requirements_text = cleaned_rfp_analysis[requirements_start:requirements_end].strip()

            mandatory_criteria = []
            for line in requirements_text.split('\n'):
                if line.strip() and ("must" in line.lower() or "required" in line.lower()):
                    # Clean each extracted criterion
                    mandatory_criteria.append(remove_problematic_chars(line.strip()))

            return mandatory_criteria
        except:
            return []

    def extract_weighted_criteria(self, rfp_analysis):
        """Extract weighted evaluation criteria from RFP analysis"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        try:
            criteria_start = cleaned_rfp_analysis.find("EVALUATION CRITERIA") + len("EVALUATION CRITERIA")
            criteria_end = cleaned_rfp_analysis.find("CLIENT PAIN POINTS", criteria_start)
            criteria_text = cleaned_rfp_analysis[criteria_start:criteria_end].strip()

            weighted_criteria = []
            for line in criteria_text.split('\n'):
                if line.strip():
                    match = re.match(r'^(.*?)(\s+\((\d+)%\))?', line.strip())
                    if match:
                        criterion = remove_problematic_chars(match.group(1).strip()) # Clean criterion text
                        weight = int(match.group(3)) if match.group(3) else 100 # Default to 100 if weight not specified
                        weighted_criteria.append((criterion, weight))
            # Default weights if none are found explicitly in RFP analysis
            if not weighted_criteria:
                 # These defaults are used if the RFP analysis doesn't explicitly list weighted criteria
                 weighted_criteria = [("Requirement Match", 40), ("Compliance", 25), ("Quality", 20), ("Alignment", 15)] # Example defaults
            return weighted_criteria
        except:
            # Return some default criteria if parsing fails
            return [("Requirement Match", 40), ("Compliance", 25), ("Quality", 20), ("Alignment", 15)]


    def extract_deadlines(self, rfp_analysis):
        """Extract deadlines from RFP analysis"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        try:
            timeline_start = cleaned_rfp_analysis.find("TIMELINE") + len("TIMELINE")
            timeline_end = cleaned_rfp_analysis.find("\n\n", timeline_start)
            timeline_text = cleaned_rfp_analysis[timeline_start:timeline_end].strip()

            deadlines = []
            for line in timeline_text.split('\n'):
                if line.strip() and any(term in line.lower() for term in ["deadline", "date", "due"]):
                    # Clean each extracted deadline
                    deadlines.append(remove_problematic_chars(line.strip()))

            return deadlines
        except:
            return []

    def extract_deliverables(self, rfp_analysis):
        """Extract deliverables from RFP analysis"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        try:
            deliverables_start = cleaned_rfp_analysis.find("DELIVERABLES") + len("DELIVERABLES")
            deliverables_end = cleaned_rfp_analysis.find("\n\n", deliverables_start)
            deliverables_text = cleaned_rfp_analysis[deliverables_start:deliverables_end].strip()

            deliverables = []
            for line in deliverables_text.split('\n'):
                if line.strip():
                    # Clean each extracted deliverable
                    deliverables.append(remove_problematic_chars(line.strip()))

            return deliverables
        except:
            return []

    def assess_compliance(self, rfp_analysis, internal_capabilities):
        """Assess compliance with internal capabilities"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        try:
            requirements_pattern = r"KEY REQUIREMENTS(.*?)DELIVERABLES"
            requirements_text = re.search(requirements_pattern, cleaned_rfp_analysis, re.DOTALL)
            requirements_text = requirements_text.group(1).strip() if requirements_text else ""

            # Ensure internal capabilities strings are cleaned
            cleaned_internal_capabilities = {
                key: [remove_problematic_chars(item) for item in value]
                for key, value in internal_capabilities.items()
            }


            prompt = f"""
            Assess compliance between RFP requirements and internal capabilities:

            RFP REQUIREMENTS:
            {requirements_text}

            INTERNAL CAPABILITIES:
            Technical: {', '.join(cleaned_internal_capabilities.get('technical', []))}
            Functional: {', '.join(cleaned_internal_capabilities.get('functional', []))}

            For each RFP requirement, determine if we can:
            - Fully comply
            - Partially comply

            Flag any requirements where we have significant gaps. Provide specific explanations for each compliance status.

            Format your response as a structured markdown table with columns:
            | Requirement | Compliance Status | Explanation |
            """

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            # Clean the generated compliance assessment text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error assessing compliance: {str(e)}")
            return "Error assessing compliance."

    def extract_required_sections(self, rfp_analysis):
        """Extract required sections from RFP analysis"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        try:
            sections_start = cleaned_rfp_analysis.find("REQUIRED SECTIONS") + len("REQUIRED SECTIONS")
            sections_end = cleaned_rfp_analysis.find("\n\n", sections_start)
            sections_text = cleaned_rfp_analysis[sections_start:sections_end].strip()
            # Clean each extracted section name
            sections = [remove_problematic_chars(s.strip()) for s in sections_text.split("\n") if s.strip()]
            return sections
        except:
            return []

    def generate_section(self, section_name, rfp_analysis, rfp_section_content,
                         client_background, differentiators,
                         evaluation_criteria, relevant_kb_content, client_name):
        """Generate a proposal section with checks for KB availability for pricing."""

        # Inputs are assumed to be cleaned by the calling function (generate_full_proposal)
        # For safety, we can re-apply cleaning here if called directly elsewhere.
        cleaned_section_name = remove_problematic_chars(section_name)
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        cleaned_rfp_section_content = remove_problematic_chars(rfp_section_content) if rfp_section_content else ""
        cleaned_client_background = remove_problematic_chars(client_background) if client_background else ""
        cleaned_differentiators = remove_problematic_chars(differentiators) if differentiators else ""
        cleaned_evaluation_criteria = remove_problematic_chars(evaluation_criteria) if evaluation_criteria else ""
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""

        # relevant_kb_content is a list of dicts; ensure content within is cleaned
        cleaned_relevant_kb_content = []
        for item in relevant_kb_content:
            if isinstance(item, dict) and 'document' in item and isinstance(item['document'], dict):
                 item['document']['filename'] = remove_problematic_chars(item['document'].get('filename', ''))
                 item['document']['section_name'] = remove_problematic_chars(item['document'].get('section_name', ''))
                 item['document']['content'] = remove_problematic_chars(item['document'].get('content', ''))
                 cleaned_relevant_kb_content.append(item)
            # else: skip malformed items


        is_pricing = any(term in cleaned_section_name.lower() for term in ["commercial", "pricing", "cost", "financial", "budget", "price"])

        pricing_block = ""
        if is_pricing:
            # --- ADDED CHECK for KB before accessing pricing ---
            if not self.kb or not hasattr(self.kb, 'extract_pricing_from_kb'):
                st.warning(f"Knowledge Base unavailable for pricing insights in section '{cleaned_section_name}'.")
                pricing_block = "\n\n## PRICING INSIGHT\nKnowledge Base unavailable."
                prices = [] # Define prices as empty list
            else:
                try:
                    prices = self.kb.extract_pricing_from_kb() # Method returns list of ints
                    if prices:
                        avg = sum(prices) / len(prices)
                        pricing_block = (
                            f"\n\n## PRICING INSIGHT\n"
                            f"Based on {len(prices)} past proposals, prices ranged from ₹{min(prices):,} "
                            f"to ₹{max(prices):,}, with an average of ₹{avg:,.0f}."
                        )
                    else:
                        pricing_block = "\n\n## PRICING INSIGHT\nNo past pricing data found in KB."
                except Exception as e:
                    st.error(f"Error extracting pricing from KB: {e}")
                    pricing_block = "\n\n## PRICING INSIGHT\nError extracting pricing data from KB."
                    prices = [] # Define prices as empty list on error
            # --- END CHECK ---
        else:
             prices = [] # Ensure prices is defined if not a pricing section

        # Prepare KB items string from the cleaned list
        kb_items = "\n\n".join([
             f"--- {('Very Relevant' if item.get('score', 0)>0.8 else 'Relevant')} PAST PROPOSAL ---\n"
             f"From: {item['document']['filename']} | Section: {item['document']['section_name']}\n"
             f"{item['document']['content']}" # Content is already cleaned
             for item in cleaned_relevant_kb_content if item.get('score', 0) >= 0.5
        ])[:2000] # Limit length

        # Ensure all parts of the prompt are cleaned strings
        prompt = f"""
        # STRATEGIC PROPOSAL SECTION GENERATION

        Section to Create: \"{cleaned_section_name}\"

        RFP CONTEXT:
        {cleaned_rfp_analysis}

        CLIENT BACKGROUND:
        {cleaned_client_background}

        EVALUATION CRITERIA:
        {cleaned_evaluation_criteria}

        DIFFERENTIATORS:
        {cleaned_differentiators}

        REFERENCE MATERIAL:
        {remove_problematic_chars(kb_items)}

        GENERATION INSTRUCTIONS:
        1. Address RFP requirements for '{cleaned_section_name}'.
        2. Use client-specific language and examples where possible, referring to '{cleaned_client_name}'.
        3. Incorporate relevant details from REFERENCE MATERIAL without direct copying.
        4. Highlight how our differentiators ({cleaned_differentiators}) meet the client's needs in this section.
        5. Ensure professional tone and clear structure.
        6. Only include explicit pricing details if this is a commercial/pricing section.
        {remove_problematic_chars(pricing_block)}
        """
        try:
            res = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"system","content":"You are an expert proposal writer, tailoring content specifically for the client and RFP section."},
                          {"role":"user","content":prompt}],
                temperature=0.2
            )
            # Clean the generated section content before returning
            return remove_problematic_chars(res.choices[0].message.content)
        except Exception as e:
            st.error(f"Error generating section '{cleaned_section_name}' via LLM: {e}")
            return f"Error generating section {cleaned_section_name}: {str(e)}" # Return cleaned error message

    def validate_proposal_client_specificity(self, proposal_sections, client_name):
        """Validates that the proposal is sufficiently client-specific"""
        issues = []
        # Ensure client name is cleaned for comparison
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""

        for section_name, content in proposal_sections.items():
            # Ensure section name and content are cleaned for validation
            cleaned_section_name = remove_problematic_chars(section_name)
            cleaned_content = remove_problematic_chars(content)

            client_name_count = cleaned_content.lower().count(cleaned_client_name.lower())
            content_length = len(cleaned_content)

            expected_mentions = max(3, content_length // 500)

            if cleaned_client_name and client_name_count < expected_mentions:
                issues.append(f"Section '{cleaned_section_name}' has insufficient client references ({client_name_count} found, {expected_mentions} expected)")

            generic_phrases = [
                "our clients", "many organizations", "typical companies",
                "best practices", "industry standards", "our approach",
                "our methodology", "our process", "our solution"
            ]

            for phrase in generic_phrases:
                if phrase in cleaned_content.lower():
                    issues.append(f"Section '{cleaned_section_name}' contains generic phrase: '{phrase}'")

        return issues

    def refine_section(self, section_name, current_content, feedback, client_name):
        """Refine a section based on user feedback"""
        # Ensure all input text is cleaned before sending to LLM
        cleaned_section_name = remove_problematic_chars(section_name)
        cleaned_current_content = remove_problematic_chars(current_content)
        cleaned_feedback = remove_problematic_chars(feedback)
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""

        # Perform replacements on cleaned content
        cleaned_current_content = cleaned_current_content.replace("CLIENT_NAME", cleaned_client_name)
        cleaned_current_content = cleaned_current_content.replace("COMPANY_NAME", "Your Company Name") # Assuming company name is safe ASCII

        prompt = f"""
        # SECTION REFINEMENT

        ## CURRENT SECTION CONTENT
        {cleaned_current_content}

        ## USER FEEDBACK
        {cleaned_feedback}

        ## REFINEMENT INSTRUCTIONS
        Revise the section to address the feedback provided. Maintain the professional tone and structure while incorporating the suggested improvements. If the feedback suggests adding specific information, ensure it's included in a relevant part of the section. If the feedback suggests restructuring, improve the organization while preserving all essential content.

        Provide the refined section content.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            # Clean the generated refined section content
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error refining section {cleaned_section_name}: {str(e)}")
            return f"Error refining section {cleaned_section_name}: {str(e)}"

    def generate_compliance_matrix(self, rfp_analysis):
        """Generate a compliance matrix using the new prompt"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        key_requirements_pattern = r"KEY REQUIREMENTS(.*?)DELIVERABLES"
        key_requirements = re.search(key_requirements_pattern, cleaned_rfp_analysis, re.DOTALL)
        key_requirements = key_requirements.group(1).strip() if key_requirements else ""

        prompt = f"""
        Create a comprehensive compliance matrix that maps RFP requirements to our proposal sections.

        Use the following RFP analysis to identify all requirements:
        {key_requirements}

        For each requirement:
        1. Quote the exact requirement language from the RFP
        2. Identify which proposal section(s) address this requirement
        3. Provide a brief (1-2 sentence) explanation of how our proposal addresses this requirement
        4. Indicate compliance status: "Fully Compliant", "Partially Compliant"

        Format the output as a structured markdown table with the following columns:
        | RFP Requirement | Reference | Addressing Section(s) | How Addressed | Compliance Status |

        Where "Reference" refers to the section/page number in the original RFP.

        Ensure every requirement from the KEY REQUIREMENTS section of the RFP analysis is included.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            # Clean the generated compliance matrix text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error generating compliance matrix: {str(e)}")
            return "Error generating compliance matrix."

    def perform_risk_assessment(self, rfp_analysis):
        """Generate a risk assessment using the new prompt"""
        # Ensure input analysis text is cleaned
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        prompt = f"""
        Create a comprehensive risk assessment for this proposal based on the following RFP analysis:
        {cleaned_rfp_analysis}

        Identify specific risks in the following categories:

        1. TECHNICAL RISKS: Integration challenges, technology limitations, compatibility issues
        2. TIMELINE RISKS: Schedule constraints, dependencies, resource availability
        3. SCOPE RISKS: Unclear requirements, potential scope changes, feature creep
        4. CLIENT RELATIONSHIP RISKS: Communication challenges, alignment issues, expectation management
        5. DELIVERY RISKS: Quality assurance, testing limitations, deployment challenges
        6. EXTERNAL RISKS: Market conditions, regulatory issues, third-party dependencies

        For each identified risk, provide:
        1. A specific, concrete description of the risk
        2. Probability assessment (Low, Medium, High) with brief justification
        3. Impact assessment (Low, Medium, High) with brief justification
        4. Specific mitigation strategy including both preventive and contingency approaches
        5. Risk owner (which team or role should manage this risk)

        Format as a well-structured markdown table. Prioritize the top 2-3 risks in each category rather than creating an exhaustive list. Focus on risks specific to this client and project.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            # Clean the generated risk assessment text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error generating risk assessment: {str(e)}")
            return "Error generating risk assessment."

    def research_client_background(self, client_name):
        """Research client background using the new prompt"""
        # Ensure client name is cleaned before sending to LLM
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""

        prompt = f"""
        Based on the client name '{cleaned_client_name}', create a strategic client profile for proposal customization.

        The profile should include:

        1. ORGANIZATION OVERVIEW:
            - Industry position and primary business focus
            - Approximate size (employees, revenue if public)
            - Geographic presence and market focus
            - Key products or services

        2. STRATEGIC PRIORITIES:
            - Current business challenges or transformation initiatives
            - Recent technology investments or digital initiatives
            - Growth areas or new market entries
            - Corporate values or mission emphasis

        3. DECISION-MAKING CONTEXT:
            - Organizational structure relevant to this proposal
            - Likely stakeholders and their priorities
            - Previous vendor relationships or relevant partnerships
            - Procurement or decision-making approach if known

        4. TECHNOLOGY LANDSCAPE:
            - Current systems or platforms likely in use
            - Technology stack preferences if known
            - Prior implementation successes or challenges
            - Digital maturity assessment

        Focus on factual information that can be verified. Where specific details aren't available, provide industry-standard insights that would still be relevant. Include only information that would directly enhance proposal customization.

        Format the response with clear headings and concise bullet points for easy reference.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.4
            )

            # Clean the generated client background text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error researching client: {str(e)}")
            return "Client background information not available."

    def evaluate_proposal_alignment(self, evaluation_criteria, proposal_sections):
        """Evaluate proposal alignment using the new prompt"""
        # Ensure input text is cleaned before sending to LLM
        cleaned_evaluation_criteria = remove_problematic_chars(evaluation_criteria)
        # Assuming proposal_sections keys are section names (already cleaned) and values are content (also cleaned)
        cleaned_proposal_sections_for_prompt = {
            remove_problematic_chars(name): remove_problematic_chars(content)
            for name, content in proposal_sections.items()
        }


        prompt = f"""
        Evaluate how effectively our proposed sections align with the evaluation criteria identified in the RFP analysis.

        RFP EVALUATION CRITERIA:
        {cleaned_evaluation_criteria}

        PROPOSED SECTIONS:
        {json.dumps(cleaned_proposal_sections_for_prompt, indent=2)} # Send cleaned sections as JSON string

        For each evaluation criterion:

        1. Identify which specific proposal section(s) address this criterion
        2. Rate our coverage as:
            - STRONG: Comprehensively addresses all aspects of the criterion
            - ADEQUATE: Addresses core requirements but could be strengthened
            - NEEDS IMPROVEMENT: Insufficient coverage or missing key elements
            - ABSENT: Not addressed in current proposal structure

        3. Provide specific, actionable recommendations to strengthen our alignment, such as:
            - Content additions or emphasis changes
            - Supporting evidence or examples to include
            - Structural improvements or section reorganization
            - Cross-references between sections to reinforce key points
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            # Clean the generated alignment assessment text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error evaluating proposal alignment: {str(e)}")
            return "Error evaluating proposal alignment with RFP criteria."

    def generate_executive_summary(self, client_background, rfp_analysis, differentiators, solution_overview, client_name):
        """Generate an executive summary using the specialized prompt"""
        # Ensure all input text is cleaned before sending to LLM
        cleaned_client_background = remove_problematic_chars(client_background) if client_background else ""
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        cleaned_differentiators = remove_problematic_chars(differentiators) if differentiators else ""
        cleaned_solution_overview = remove_problematic_chars(solution_overview) if solution_overview else ""
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""


        prompt = f"""
        Create a compelling Executive Summary for this proposal based on the following inputs:

        CLIENT BACKGROUND:
        {cleaned_client_background}

        RFP ANALYSIS:
        {cleaned_rfp_analysis}

        KEY DIFFERENTIATORS:
        {cleaned_differentiators}

        SOLUTION OVERVIEW:
        {cleaned_solution_overview}

        Your Executive Summary should:

        1. Open with a concise statement acknowledging the client's specific needs and challenges
        2. Present our understanding of their primary objectives in pursuing this project
        3. Outline our proposed approach at a high level (without technical detail)
        4. Highlight 2-3 key differentiators that make our solution uniquely valuable
        5. Reference our relevant experience and qualifications specifically relevant to their needs
        6. Close with a compelling value proposition that addresses their business outcomes

        Keep the Executive Summary to approximately 500 words using clear, confident language that demonstrates both understanding and expertise. Avoid generic claims and focus on client-specific value. Use minimal formatting - short paragraphs with occasional bold text for emphasis.

        The Executive Summary should stand alone if separated from the full proposal while compelling the reader to continue.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.4
            )

            # Clean the generated executive summary text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error generating executive summary: {str(e)}")
            return f"Error generating executive summary: {str(e)}"

    def generate_full_proposal(self, rfp_text, client_name=None, company_info=None, template_sections=None):
        """Generate a full proposal with checks for KB initialization."""

        # --- ADDED CHECK ---
        # Check if the Knowledge Base is initialized and has the required methods
        if not self.kb or not hasattr(self.kb, 'multi_hop_search') or not hasattr(self.kb, 'extract_pricing_from_kb'):
            st.error("Knowledge Base is not properly initialized within the Proposal Generator. Cannot generate full proposal.")
            # Return an error structure consistent with the expected output
            return {
                "analysis": "Error: Knowledge Base not initialized.",
                "sections": {},
                "client_background": "Client background not available due to KB error.",
                "differentiators": company_info.get("differentiators", "") if company_info else "",
                "required_sections": template_sections or [],
                "client_name": remove_problematic_chars(client_name) if client_name else None
            }
        # --- END CHECK ---

        print("Analyzing RFP...")
        # Clean RFP text before analysis
        cleaned_rfp_text = remove_problematic_chars(rfp_text)
        rfp_analysis = self.analyze_rfp(cleaned_rfp_text) # Analysis result is cleaned by the method

        if template_sections:
            # Ensure template sections are cleaned
            required_sections = [remove_problematic_chars(s) for s in template_sections]
        else:
            # extract_required_sections uses cleaned analysis and returns cleaned sections
            required_sections = self.extract_required_sections(rfp_analysis)
            if not required_sections: # Fallback if LLM fails extraction or returns empty
                rfp_doc_sections = extract_sections_from_rfp(cleaned_rfp_text)
                required_sections = [remove_problematic_chars(s) for s in rfp_doc_sections.keys()] if rfp_doc_sections else ["Introduction", "Proposed Solution", "Pricing", "Conclusion"]
                st.warning(f"Could not extract specific required sections from RFP Analysis. Using sections: {', '.join(required_sections)}")


        # Clean client name and company info before research/use
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else None
        cleaned_company_info = {
            "name": remove_problematic_chars(company_info.get("name", "")) if company_info else "",
            "differentiators": remove_problematic_chars(company_info.get("differentiators", "")) if company_info else ""
        } if company_info else {}


        if cleaned_client_name:
            # research_client_background returns cleaned background
            client_background = self.research_client_background(cleaned_client_name)
        else:
            client_background = "Client background not provided."

        differentiators = cleaned_company_info.get("differentiators", "Company differentiators not provided.")


        criteria_pattern = r"EVALUATION CRITERIA(.*?)CLIENT PAIN POINTS"
        # Use cleaned rfp_analysis for pattern matching
        evaluation_criteria_match = re.search(criteria_pattern, rfp_analysis, re.DOTALL)
        # evaluation_criteria is cleaned after extraction
        evaluation_criteria = remove_problematic_chars(evaluation_criteria_match.group(1).strip()) if evaluation_criteria_match else "Evaluation criteria not specified."

        proposal_sections = {}
        # Extract sections from the cleaned RFP text once before the loop
        rfp_sections_content_map = extract_sections_from_rfp(cleaned_rfp_text)

        for section_name in required_sections: # required_sections are already cleaned
            print(f"Generating section: {section_name}")

            # Find corresponding RFP section content (case-insensitive matching)
            rfp_section_content_for_llm = next((content for rfp_sec_name, content in rfp_sections_content_map.items() if section_name.lower() in rfp_sec_name.lower() or rfp_sec_name.lower() in section_name.lower()), "")
            # Content is already cleaned by extract_sections_from_rfp

            cleaned_rfp_section_content = remove_problematic_chars(rfp_section_content_for_llm) if rfp_section_content_for_llm else ""
            expanded_query = expand_query(section_name + " " + cleaned_rfp_section_content)

            # --- ADDED TRY-EXCEPT around KB search ---
            relevant_kb_content = [] # Default to empty list
            try:
                # Assuming self.kb was validated at the start of the method
                relevant_kb_content = self.kb.multi_hop_search(expanded_query, k=3) # multi_hop_search returns cleaned content
            except Exception as kb_error:
                st.error(f"Error searching Knowledge Base for section '{section_name}': {kb_error}")
                # Continue generation with empty KB content
            # --- END TRY-EXCEPT ---

            # Call generate_section (which now also has KB checks for pricing)
            # All inputs passed here should be cleaned versions
            proposal_sections[section_name] = self.generate_section(
                section_name,           # Cleaned
                rfp_analysis,           # Cleaned
                cleaned_rfp_section_content, # Cleaned
                client_background,      # Cleaned
                differentiators,        # Cleaned
                evaluation_criteria,    # Cleaned
                relevant_kb_content,    # Contains cleaned content
                cleaned_client_name     # Cleaned
            )

        # Generate Executive Summary if needed
        # Check against cleaned section names in the generated proposal_sections dictionary
        if "Executive Summary" not in proposal_sections and cleaned_client_name:
            print("Generating Executive Summary...")

            section_highlights = ""
            key_sections_for_summary = ["Approach", "Methodology", "Solution", "Benefits", "Implementation"]
            for summary_key_sec in key_sections_for_summary:
                # Match cleaned section names from the generated proposal
                matching_gen_section = next((s_name for s_name in proposal_sections.keys() if summary_key_sec.lower() in s_name.lower()), None)
                if matching_gen_section:
                    # Use cleaned proposal section content for highlights
                    content_preview = remove_problematic_chars(proposal_sections[matching_gen_section])[:200] + "..."
                    section_highlights += f"## {matching_gen_section} Preview\n{content_preview}\n\n"

            # Ensure section_highlights is cleaned (though composed from cleaned parts)
            cleaned_section_highlights = remove_problematic_chars(section_highlights)

            # Generate the executive summary using cleaned inputs
            try:
                # generate_executive_summary handles cleaning internally now
                exec_summary_content = self.generate_executive_summary(
                     client_background,     # Cleaned
                     rfp_analysis,          # Cleaned
                     differentiators,       # Cleaned
                     cleaned_section_highlights, # Cleaned overview
                     cleaned_client_name    # Cleaned
                )
                proposal_sections["Executive Summary"] = exec_summary_content # Result is cleaned by generate_executive_summary
            except Exception as e:
                 print(f"Error generating Executive Summary: {str(e)}")
                 proposal_sections["Executive Summary"] = "Error generating Executive Summary." # Add placeholder on error

        # Final structure uses cleaned data
        return {
            "analysis": rfp_analysis,
            "sections": proposal_sections,
            "client_background": client_background,
            "differentiators": differentiators,
            "required_sections": required_sections,
            "client_name": cleaned_client_name
        }


    def perform_quality_assurance(self, proposal_sections, rfp_analysis):
        """Perform quality assurance checks on the proposal"""
        # Clean proposal sections before sending to LLM for QA
        cleaned_proposal_sections = {remove_problematic_chars(name): remove_problematic_chars(content) for name, content in proposal_sections.items()}

        prompt = f"""
        # QUALITY ASSURANCE CHECK

        ## PROPOSAL SECTIONS
        {json.dumps(cleaned_proposal_sections, indent=2)}

        ## RFP ANALYSIS
        {remove_problematic_chars(rfp_analysis)} # Clean RFP analysis too

        ## QUALITY ASSURANCE INSTRUCTIONS
        Perform comprehensive quality assurance checks on the proposal sections:

        1. LANGUAGE TONE:
            - Evaluate if the tone is professional and confident
            - Check for overly technical language that might confuse the client
            - Identify any overly casual or informal language

        2. GRAMMAR AND SPELLING:
            - Identify any grammatical errors
            - Check for spelling mistakes
            - Verify consistency in verb tenses and subject-verb agreement

        3. COMPLIANCE WITH RFP GUIDELINES:
            - Verify that all sections address the RFP requirements
            - Check if the proposal follows the structure requested in the RFP
            - Ensure all mandatory elements from the RFP are included

        4. CONTENT QUALITY:
            - Evaluate if claims are supported with evidence
            - Check for vague statements that should be more specific
            - Identify any sections that could benefit from additional details

        Provide specific feedback for each section, including:
        - Exact location of the issue
        - Description of the problem
        - Suggested improvement

        Format the output as a structured markdown report with the following sections:
        ## Overall Quality Score (1-10)
        ## Tone Assessment
        ## Grammar and Spelling Issues
        ## Compliance with RFP Guidelines
        ## Content Quality Feedback
        ## Actionable Improvements
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            # Clean the generated QA text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error performing quality assurance: {str(e)}")
            return "Error performing quality assurance."

    def generate_advanced_analysis(self, proposal_data, rfp_analysis, internal_capabilities, client_name):
        """Generate advanced analysis without executive summary"""
        analysis_results = {}

        # Clean inputs before passing to generation functions
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        # Clean proposal section names (keys) for sending to LLM
        cleaned_proposal_sections_keys = [remove_problematic_chars(name) for name in proposal_data["sections"].keys()]
        cleaned_internal_capabilities = {
            key: [remove_problematic_chars(item) for item in value]
            for key, value in internal_capabilities.items()
        }
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""

        compliance_matrix = self.generate_compliance_matrix(cleaned_rfp_analysis)
        analysis_results["compliance_matrix"] = compliance_matrix

        risk_assessment = self.perform_risk_assessment(cleaned_rfp_analysis)
        analysis_results["risk_assessment"] = risk_assessment

        criteria_pattern = r"EVALUATION CRITERIA(.*?)CLIENT PAIN POINTS"
        evaluation_criteria = re.search(criteria_pattern, cleaned_rfp_analysis, re.DOTALL)
        evaluation_criteria = evaluation_criteria.group(1).strip() if evaluation_criteria else "Evaluation criteria not specified."
        cleaned_evaluation_criteria = remove_problematic_chars(evaluation_criteria)


        alignment_assessment = self.evaluate_proposal_alignment(
            cleaned_evaluation_criteria,
            # Pass the dictionary of cleaned section names and content for alignment evaluation
            {remove_problematic_chars(name): remove_problematic_chars(content) for name, content in proposal_data["sections"].items()}
        )
        analysis_results["alignment_assessment"] = alignment_assessment

        compliance_assessment = self.assess_compliance(cleaned_rfp_analysis, cleaned_internal_capabilities)
        analysis_results["compliance_assessment"] = compliance_assessment

        return analysis_results

    def analyze_vendor_proposal(self, vendor_proposal_text, rfp_analysis, client_name, scoring_system):
        """Analyze vendor proposal against RFP requirements with detailed factual comparison"""
        # Clean input texts before analysis
        cleaned_vendor_proposal_text = remove_problematic_chars(vendor_proposal_text)
        cleaned_rfp_analysis = remove_problematic_chars(rfp_analysis)
        cleaned_client_name = remove_problematic_chars(client_name) if client_name else ""

        # Extract specific RFP requirements for comparison
        weighted_criteria = self.extract_weighted_criteria(cleaned_rfp_analysis)

         # Format scoring metrics and weights from config for the prompt
        # Use the metrics from the scoring_system passed in (which could be dynamic)
        scoring_metrics_info = "\n".join([f"- {metric.replace('_', ' ').title()}" # Only include metric name in prompt, not dynamic weight
                                            for metric in scoring_system['weighting'].keys()])

        # Generate analysis prompt with detailed instructions
        analysis_prompt = f"""
        # DETAILED VENDOR PROPOSAL ANALYSIS

        ## ANALYSIS INSTRUCTIONS:
        1. Perform a comprehensive analysis of the vendor proposal against the provided RFP requirements.
        2. Evaluate the proposal based on the defined scoring metrics.
        3. For each scoring metric listed below, provide a score out of 100 based on your detailed analysis.
        4. Provide specific examples from both the RFP and proposal to support your analysis.
        5. Format your response with clear headings for each analysis category and clearly state the score for each metric using the format "**[Metric Name] Score: [Score]/100**".

        ## SCORING METRICS TO EVALUATE (Score each out of 100):
        {scoring_metrics_info}


        ## RFP REQUIREMENTS:
        {cleaned_rfp_analysis} # Use cleaned RFP analysis

        ## VENDOR PROPOSAL:
        {cleaned_vendor_proposal_text} # Use cleaned vendor text

         ## ANALYSIS FORMAT:
        ### Overall Score (0-100)
        ### Requirement Matching:
        - Fully Addressed Requirements
        - Partially Addressed Requirements
        - Unaddressed Requirements
        ### Compliance Assessment:
        - Met Requirements
        - Partially Met Requirements
        - Unmet Requirements
        ### Quality Evaluation:
        - Strengths
        - Weaknesses
        ### Alignment with Client Priorities:
        - Well-Aligned Aspects
        - Misaligned Aspects
        ### Risk Assessment:
        - Identified Risks
        - Mitigation Strategies
        ### Sentiment Analysis:
        - Tone assessment (positive, neutral, negative)
        - Confidence level of the vendor
        ### Comparative Analysis:
        - How this proposal compares to typical industry standards
        - Competitive advantages/disadvantages

        Provide specific page/section references from the proposal for each assessment point.
        Calculate an overall score based on the weighted metrics.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini", # Consider GPT-4 for potentially better scoring consistency
                messages=[
                    {"role": "system", "content": "You are an expert proposal evaluator providing detailed analysis and scoring."},
                    {"role": "user", "content": analysis_prompt}
                ],
                temperature=0.1 # Lower temperature for more factual and consistent scoring
            )
            analysis_text = response.choices[0].message.content

            # Clean the generated analysis text
            cleaned_analysis_text = remove_problematic_chars(analysis_text)

            return cleaned_analysis_text
        except Exception as e:
            print(f"Error analyzing vendor proposal: {str(e)}")
            return f"Error analyzing vendor proposal: {str(e)}\n\nPrompt:\n{analysis_prompt}" # Return prompt on error for debugging


    def calculate_weighted_score(self, analysis_text: str, scoring_system: Dict) -> Tuple[Optional[float], Dict[str, Optional[int]], Optional[str]]:
        """
        Parses vendor analysis text to extract scores for configured metrics,
        calculates the weighted score, and determines the grade.
        Updated regex to be more flexible with LLM output variations.

        Args:
            analysis_text: The text output from analyze_vendor_proposal.
            scoring_system: The scoring configuration dictionary, potentially with dynamic weights.

        Returns:
            A tuple containing:
            - The total weighted score (float) or None if calculation fails.
            - A dictionary of individual metric scores (str: int) or None if not found.
            - The calculated grade (str) based on the score, or None.
        """
        # Ensure analysis text is cleaned before parsing
        cleaned_analysis_text = remove_problematic_chars(analysis_text)

        weights = scoring_system.get('weighting', {})
        grading_scale = scoring_system.get('grading_scale', {})
        individual_scores = {}
        total_weighted_score = 0.0
        total_weight_sum = sum(weights.values()) # Calculate the sum of weights provided

        # Define regex patterns based on the metrics provided in the scoring_system
        for metric in weights.keys(): # Use the keys from the provided weights
            # Create a user-friendly metric name for the regex (e.g., "requirement_match" -> "Requirement Match")
            metric_name_formatted = metric.replace('_', ' ').title()
            # Updated regex: Look for the formatted metric name, followed by anything (non-greedy),
            # then capture either digits (\d+) or "N/A". This is more robust to variations
            # in the LLM's output format after the metric name.
            pattern = re.compile(rf"{re.escape(metric_name_formatted)}.*?(\d+|N/A)", re.IGNORECASE | re.DOTALL)
            match = pattern.search(cleaned_analysis_text) # Search in cleaned text

            if match:
                score_str = match.group(1)
                if score_str.upper() == 'N/A':
                    score = None # Represent N/A as None
                else:
                    try:
                        score = int(score_str)
                        # Ensure score is within 0-100 range if it's a number
                        score = max(0, min(100, score))
                    except ValueError:
                        score = None # Handle cases where captured text isn't a valid integer

                individual_scores[metric] = score

                # Only include valid numerical scores in the weighted calculation
                if score is not None:
                    # Multiply the individual score (out of 100) by its weight
                    total_weighted_score += score * weights[metric]
                print(f"Found score for {metric}: {score_str} -> {score}, weight: {weights.get(metric, 'N/A')}") # Debug print
            else:
                individual_scores[metric] = None
                print(f"Could not find score for {metric}") # Debug print

        # Normalize the total weighted score if the total weight sum is not 1 or 100
        # If weights sum to 100, this effectively scales the score to 0-100
        # If weights sum to 1, the score is already out of 100 conceptually
        # If weights sum to something else, normalize by the sum.
        final_score_raw = total_weighted_score
        # Only normalize if there's a positive total weight
        final_score_for_grading = (final_score_raw / total_weight_sum) if total_weight_sum > 0 else 0
        # Cap the score at 100 for grading purposes, even if raw calculation exceeds it
        final_score_for_grading = min(final_score_for_grading, 100)
        # Ensure score is not negative
        final_score_for_grading = max(0, final_score_for_grading)


        print(f"Raw Weighted Score: {final_score_raw}, Score for Grading (Normalized): {final_score_for_grading:.2f}, Total Weight Sum: {total_weight_sum}") # Debug print

        # Determine grade
        grade = "N/A"
        # Sort grading scale by the lower bound in descending order to ensure correct grade assignment
        sorted_grading_scale = sorted(grading_scale.items(), key=lambda item: item[1][0], reverse=True)
        for grade_name, score_range in sorted_grading_scale:
            # Ensure score_range has two elements and they are numbers
            if isinstance(score_range, list) and len(score_range) == 2 and all(isinstance(s, (int, float)) for s in score_range):
                if score_range[0] <= final_score_for_grading <= score_range[1]:
                    grade = grade_name.title()
                    break
            else:
                print(f"Warning: Invalid grading scale format for '{grade_name}': {score_range}") # Debug print


        print(f"Calculated Grade: {grade}") # Debug print

        # Return the normalized score for display
        return final_score_for_grading, individual_scores, grade


    def identify_gaps_and_risks(self, vendor_proposal_text, rfp_requirements):
        """Use machine learning to identify gaps and risks in vendor responses"""
        # Clean input texts before processing
        cleaned_vendor_proposal_text = remove_problematic_chars(vendor_proposal_text)
        cleaned_rfp_requirements = remove_problematic_chars(rfp_requirements)

        try:
            # Vectorize text
            documents = [cleaned_rfp_requirements, cleaned_vendor_proposal_text]
            vectorizer = TfidfVectorizer()
            tfidf_matrix = vectorizer.fit_transform(documents)

            # Calculate similarity
            similarity_matrix = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
            similarity_score = similarity_matrix[0][0] # Get the single similarity value

            gaps = []
            if similarity_score < 0.7: # Threshold for identifying gaps
                gaps.append(f"Potential low coverage of key requirements (Similarity Score: {similarity_score:.2f})")
            if similarity_score < 0.5:
                gaps.append("Potential mismatch in proposed solutions compared to requirements")

            # Basic keyword-based risk identification (can be expanded)
            risks = []
            risk_keywords = ["unable to", "cannot commit", "significant challenge", "out of scope", "additional cost", "dependency on client"]
            # Iterate through the vendor proposal text to find occurrences of risk keywords
            # This is a simple approach and could be improved with NLP techniques
            cleaned_vendor_text_lower = cleaned_vendor_proposal_text.lower() # Use cleaned text for keyword search
            for keyword in risk_keywords:
                # Use re.findall to find all occurrences, case-insensitive
                if re.search(r'\b' + re.escape(keyword) + r'\b', cleaned_vendor_text_lower):
                    risks.append(f"Potential risk identified related to keyword: '{keyword}'")

            if similarity_score < 0.5:
                risks.append("High risk of non-compliance due to low overall similarity")

            # Ensure extracted gaps and risks strings are cleaned
            cleaned_gaps = [remove_problematic_chars(g) for g in gaps]
            cleaned_risks = [remove_problematic_chars(r) for r in risks]

            return cleaned_gaps, cleaned_risks
        except Exception as e:
            print(f"Error identifying gaps and risks: {str(e)}")
            return [], []

    def generate_scoring_analysis(self, vendor_analyses):
        """Generate comprehensive scoring analysis for multiple vendor proposals"""
        scores = []
        # Clean vendor analyses text before processing
        cleaned_vendor_analyses = [remove_problematic_chars(analysis) for analysis in vendor_analyses]

        for analysis in cleaned_vendor_analyses:
            try:
                # This method is currently not used in the UI tabs provided.
                # If it were, it would need to be updated to handle potentially
                # different scoring metrics/weights per vendor analysis if
                # analyses were done with different configurations.
                # For now, keeping it as is based on the original code.
                score_match = re.search(r"match\s*score\s*:\s*(\d+)", analysis, re.IGNORECASE)
                if score_match:
                    score = int(score_match.group(1))
                    scores.append(score)
            except:
                scores.append(0)

        if not scores:
            return "No scores found for analysis."

        # Calculate statistics
        avg_score = sum(scores) / len(scores)
        max_score = max(scores) if scores else 0
        min_score = min(scores) if scores else 0

        # Generate analysis prompt
        prompt = f"""
        # VENDOR PROPOSAL SCORING ANALYSIS

        Analyze the following scores from vendor proposals:

        Scores: {', '.join(map(str, scores))}

        Calculate:
        - Average Score: {avg_score:.1f}
        - Maximum Score: {max_score}
        - Minimum Score: {min_score}

        Provide insights into:
        - How vendors performed against each other
        - Which vendors demonstrated superior alignment
        - Common strengths across proposals
        - Recurring weaknesses or gaps
        - Recommendations for vendor selection based on scores

        Format your response with clear headings and bullet points.
        """

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )

            # Clean the generated scoring analysis text
            return remove_problematic_chars(response.choices[0].message.content)
        except Exception as e:
            print(f"Error generating scoring analysis: {str(e)}")
            return f"Error generating scoring analysis: {str(e)}"

# Word export function
def export_to_word(proposal_data, company_name, client_name, output_path, company_logo_path=None):
    """Export the generated proposal to a professionally formatted Word document"""
    doc = Document()

    # Set document styles
    styles = doc.styles

    # Modify heading styles
    heading1 = styles['Heading 1']
    heading1.font.size = Pt(16)
    heading1.font.bold = True

    heading2 = styles['Heading 2']
    heading2.font.size = Pt(14)
    heading2.font.bold = True

    # Set document properties
    doc.core_properties.author = remove_problematic_chars(company_name) if company_name else ""
    doc.core_properties.title = remove_problematic_chars(f"Proposal for {client_name}") if client_name else "Proposal"


    # Add title page
    if company_logo_path and os.path.exists(company_logo_path):
        try:
            doc.add_picture(company_logo_path, width=Inches(2.0))
            doc.add_paragraph()  # Add some space after logo
        except Exception as e:
            print(f"Could not add logo to document: {e}")


    # Ensure client_name is cleaned before adding to document
    cleaned_client_name = remove_problematic_chars(client_name) if client_name else "Client"
    title = doc.add_heading(f"Proposal for {cleaned_client_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph()
    # Ensure company_name is cleaned before adding to document
    cleaned_company_name = remove_problematic_chars(company_name) if company_name else "Your Company Name"
    subtitle_run = subtitle.add_run(f"Prepared by {cleaned_company_name}")
    subtitle_run.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a page break
    doc.add_page_break()

    # Add table of contents title
    doc.add_heading("Table of Contents", 1)

    # Generate table of contents
    # Note: Accurate TOC page numbers in docx are complex and usually require a second pass
    # or using built-in docx features which are harder to control programmatically.
    # This provides a basic list of sections.
    toc = doc.add_paragraph()
    for idx, section_name in enumerate(proposal_data["sections"]):
        # Ensure section_name is cleaned for TOC
        cleaned_section_name = remove_problematic_chars(section_name)
        toc.add_run(f"{cleaned_section_name}").bold = True
        toc.add_run(f"... [Page Number]\n") # Placeholder


    # Add page break after TOC
    doc.add_page_break()

    # Add each section with proper formatting
    for section_name, section_content in proposal_data["sections"].items():
        # Ensure section_name is cleaned for heading
        cleaned_section_name = remove_problematic_chars(section_name)
        doc.add_heading(cleaned_section_name, 1)

        # Ensure section_content is cleaned before processing lines
        cleaned_section_content = remove_problematic_chars(section_content)
        lines = cleaned_section_content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if line.startswith('### '):
                # Ensure heading text is cleaned
                doc.add_heading(remove_problematic_chars(line[4:].strip()), 3)
            elif line.startswith('## '):
                # Ensure heading text is cleaned
                doc.add_heading(remove_problematic_chars(line[3:].strip()), 2)
            elif line.startswith('# '):
                # Ensure heading text is cleaned
                doc.add_heading(remove_problematic_chars(line[2:].strip()), 1)
            elif line.startswith('- ') or line.startswith('* '):
                # Ensure list item text is cleaned
                p = doc.add_paragraph(remove_problematic_chars(line[2:]), style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                # Ensure list item text is cleaned
                p = doc.add_paragraph(remove_problematic_chars(re.sub(r'^\d+\.\s', '', line)), style='List Number')
            elif line.startswith('|') and i+1 < len(lines) and '|--' in lines[i+1]:
                # Basic table parsing
                table_rows = []
                table_rows.append(line)
                i += 1
                while i < len(lines) and lines[i].startswith('|'):
                    table_rows.append(lines[i])
                    i += 1
                if len(table_rows) > 1: # Need at least header and one data row (or just header if parsing allows)
                    # Assuming header is the first row and separator is the second
                    header_cells = [remove_problematic_chars(cell.strip()) for cell in table_rows[0].split('|')[1:-1]] # Clean header cells
                    num_cols = len(header_cells)
                    if num_cols > 0:
                        # Count data rows (excluding header and separator)
                        data_rows = [row for row in table_rows[2:] if row.strip() and '|' in row]
                        num_rows = len(data_rows) + 1 # Add 1 for the header row

                        if num_rows > 0:
                             table = doc.add_table(rows=num_rows, cols=num_cols)
                             table.style = 'Table Grid'

                             # Add header row
                             for j, cell_text in enumerate(header_cells):
                                 table.cell(0, j).text = cell_text # Header cells are already cleaned

                             # Add data rows
                             for row_idx, row_text in enumerate(data_rows):
                                 cells = [remove_problematic_chars(cell.strip()) for cell in row_text.split('|')[1:-1]] # Clean data cells
                                 for j, cell_text in enumerate(cells):
                                     if j < num_cols: # Ensure we don't go out of bounds
                                         table.cell(row_idx+1, j).text = cell_text # Data cells are already cleaned
                        else:
                             # Handle case with only a header and separator, no data rows
                             table = doc.add_table(rows=1, cols=num_cols)
                             table.style = 'Table Grid'
                             for j, cell_text in enumerate(header_cells):
                                 table.cell(0, j).text = cell_text # Header cells are already cleaned

                i -= 1 # Decrement i because the while loop incremented it past the table
            elif line:
                # Ensure paragraph text is cleaned
                p = doc.add_paragraph(remove_problematic_chars(line))
            i += 1

        if section_name != list(proposal_data["sections"].keys())[-1]:
            doc.add_page_break()

    doc.save(output_path)

    return output_path

# PDF export function
def export_to_pdf(proposal_data, company_name, client_name, output_path, company_logo_path=None):
    try:
        from fpdf import FPDF
    except ImportError:
        st.error("The 'fpdf' library is required for PDF export. Please install it: pip install fpdf")
        return None

    class ProposalPDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 12)
            # Ensure client_name is cleaned before adding to header
            client_name_str = remove_problematic_chars(str(client_name)) if client_name else "Client"
            self.cell(0, 10, txt=f"Proposal for {client_name_str}", border=0, ln=1, align='C')

            if company_logo_path and os.path.exists(company_logo_path):
                try:
                    self.image(company_logo_path, 10, 8, 30)
                except Exception as e:
                    print(f"Could not add logo to PDF: {e}")

            self.ln(20)

        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, txt=f'Page {self.page_no()}', border=0, ln=0, align='C')

    pdf = ProposalPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", size=12)

    # Table of Contents
    pdf.cell(200, 10, txt="Table of Contents", ln=True, align='C')
    pdf.ln(5)

    # Note: Calculating accurate page numbers in PDF based on content length is complex.
    # This provides a basic estimate.
    current_page = 3 # Start after Title and TOC pages
    for section_name in proposal_data["sections"]:
         # Ensure section_name is cleaned for TOC
         cleaned_section_name = remove_problematic_chars(section_name)
         pdf.cell(0, 10, txt=f"{cleaned_section_name} - Page {current_page}", ln=True)
         # Estimate pages for the next section - a rough estimate
         # This is highly dependent on font size, line height, page margins, etc.
         # A more accurate method would involve rendering the content and counting pages.
         # Ensure content used for estimation is cleaned
         content = remove_problematic_chars(proposal_data["sections"][section_name])
         lines_per_page_estimate = 40 # Rough estimate
         estimated_lines = len(content.split('\n'))
         estimated_pages = max(1, estimated_lines // lines_per_page_estimate)
         current_page += estimated_pages


    pdf.add_page()

    for section_name, content in proposal_data["sections"].items():
        # Ensure section_name is cleaned for heading
        cleaned_section_name = remove_problematic_chars(section_name)
        pdf.set_font("Arial", 'B', 16)
        pdf.multi_cell(0, 10, txt=cleaned_section_name, border=0) # Use multi_cell for long titles
        pdf.ln(5) # Reduced line break after section title

        pdf.set_font("Arial", size=12)
        # Ensure content is cleaned before splitting into lines
        cleaned_content = remove_problematic_chars(content)
        lines = cleaned_content.split('\n')
        for line in lines:
            # Handle basic markdown like bold/italic if needed, fpdf requires specific commands
            # For simplicity here, just print lines. More complex formatting requires parsing markdown.
            # Ensure line is cleaned before printing
            cleaned_line = remove_problematic_chars(line)
            cleaned_line = re.sub(r'[\*_`]', '', cleaned_line) # Remove basic markdown chars
            if cleaned_line.strip(): # Avoid adding empty lines
                 pdf.multi_cell(0, 6, txt=cleaned_line, border=0)
                 pdf.ln(1) # Reduced line break between paragraphs

        # Add page break if it's not the last section
        if section_name != list(proposal_data["sections"].keys())[-1]:
            pdf.add_page()

    pdf.output(output_path)
    return output_path


# Main Streamlit UI
def main():
    st.set_page_config(page_title="AI Proposal & RFP Generator", layout="wide", page_icon="📄")

    # Apply custom CSS
    st.markdown("""
    <style>
    .main {
        background-color: #f5f7f9;
    }
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    h1, h2, h3 {
        color: #003366;
    }
    .stButton button {
        background-color: #003366;
        color: white;
    }
    .info-box {
        background-color: #e8f4f8;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 15px;
    }
    .section-card {
        background-color: white;
        padding: 20px;
        border-radius: 5px;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        margin-bottom: 15px;
    }
    .sidebar-content {
        background-color: white;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 15px;
    }
    </style>
    """, unsafe_allow_html=True)

    # Initialize session state variables
    if 'config' not in st.session_state:
        st.session_state.config = load_config()

    # Ensure scoring_system exists in config after loading
    if 'scoring_system' not in st.session_state.config:
        st.session_state.config['scoring_system'] = {
            "weighting": {
                "requirement_match": 0.4,
                "compliance": 0.25,
                "quality": 0.2,
                "alignment": 0.15,
                "risk": 0.1
            },
            "grading_scale": {
                "excellent": [90, 100],
                "good": [70, 89],
                "fair": [50, 69],
                "poor": [0, 49]
            }
        }
        print("Added default scoring_system to config in session state.")


    if 'knowledge_base' not in st.session_state:
        try:
            kb_dir = st.session_state.config["knowledge_base"]["directory"]
            embedding_model = st.session_state.config["knowledge_base"]["embedding_model"]
            st.session_state.knowledge_base = ProposalKnowledgeBase(kb_dir, embedding_model)
        except Exception as e:
            st.error(f"Failed to initialize knowledge base: {str(e)}")
            st.session_state.knowledge_base = None

    if 'generator' not in st.session_state:
        openai_key = st.session_state.config["api_keys"]["openai_key"]
        if not openai_key:
            openai_key = os.environ.get("OPENAI_API_KEY", "")

        if openai_key:
            st.session_state.generator = EnhancedProposalGenerator(st.session_state.knowledge_base, openai_key)
        else:
            st.error("OpenAI API key is not configured. Please add it to config.json or set the OPENAI_API_KEY environment variable.")
            st.session_state.generator = None


    if 'rfp_text' not in st.session_state:
        st.session_state.rfp_text = ""

    if 'rfp_analysis' not in st.session_state:
        st.session_state.rfp_analysis = None

    if 'proposal_data' not in st.session_state:
        st.session_state.proposal_data = {
            "sections": {},
            "required_sections": [],
            "client_background": None,
            "differentiators": None,
            "client_name": "Client Organization" # Initialize client_name
        }

    if 'client_background' not in st.session_state:
        st.session_state.client_background = None

    if 'differentiators' not in st.session_state:
        st.session_state.differentiators = None

    if 'advanced_analysis' not in st.session_state:
        st.session_state.advanced_analysis = {
            "compliance_matrix": None,
            "risk_assessment": None,
            "alignment_assessment": None,
            "compliance_assessment": None
        }

    if 'template_created' not in st.session_state:
        st.session_state.template_created = False

    if 'template_sections' not in st.session_state:
        st.session_state.template_sections = []

    if 'rfp_response_analysis' not in st.session_state:
        st.session_state.rfp_response_analysis = None

    if 'vendor_analysis' not in st.session_state:
        st.session_state.vendor_analysis = None

    if 'vendor_score_results' not in st.session_state:
        st.session_state.vendor_score_results = None

    if 'vendor_gaps_risks' not in st.session_state:
        st.session_state.vendor_gaps_risks = None

    if 'vendor_proposals' not in st.session_state:
        st.session_state.vendor_proposals = []

    if 'rfp_templates' not in st.session_state:
        st.session_state.rfp_templates = []

    if 'rfp_template_content' not in st.session_state:
        st.session_state.rfp_template_content = None

    # Initialize session state for dynamic weights if not already present
    # Use the default weights from the loaded config
    # Ensure scoring_system and weighting keys exist before accessing
    if 'dynamic_weights' not in st.session_state or not st.session_state.dynamic_weights:
         # Use a copy to avoid modifying the original config in session state directly
         st.session_state.dynamic_weights = st.session_state.config.get('scoring_system', {}).get('weighting', {}).copy()
         if not st.session_state.dynamic_weights:
              # Fallback if weighting is also missing
              st.session_state.dynamic_weights = {
                 "requirement_match": 0.4,
                 "compliance": 0.25,
                 "quality": 0.2,
                 "alignment": 0.15,
                 "risk": 0.1
              }
              st.warning("Scoring system weights not found in config. Using default weights.")


    # Header
    st.title("🚀 AI-Powered Proposal & RFP Generator")

    # Main workflow tabs
    tabs = st.tabs(["📋 Upload RFP", "📝 Proposal Template Creation", "📊 Generate Proposal", "📤 Export", "🔍 Advanced Analysis", "🔍 Vendor Proposal Evaluation", "📋 RFP Template Creator"])

    # Tab 1: Upload RFP
    with tabs[0]:
        st.header("Upload and Analyze RFP")

        col1, col2 = st.columns([3, 2])

        with col1:
            uploaded_file = st.file_uploader("Upload RFP Document", type=["docx", "pdf", "txt", "md"])

            if uploaded_file is not None:
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}")
                temp_file.write(uploaded_file.getvalue())
                temp_file.close()

                try:
                    rfp_text = process_rfp(temp_file.name)
                    st.session_state.rfp_text = rfp_text
                    st.success(f"Successfully processed {uploaded_file.name}")

                    with st.expander("Preview RFP Content", expanded=False):
                        # Display cleaned text in the preview
                        st.text_area("RFP Text", remove_problematic_chars(rfp_text), height=300, key="rfp_preview")

                except Exception as e:
                    st.error(f"Error processing file: {str(e)}")
                finally:
                    if os.path.exists(temp_file.name):
                        os.unlink(temp_file.name)

        with col2:
            st.markdown('<div class="info-box">', unsafe_allow_html=True)
            st.markdown("### 📝 Instructions")
            st.markdown("""
            1. Upload your RFP document (PDF, Word, or text)
            2. We'll extract and analyze the key requirements
            3. Click 'Analyze RFP' to get insights
            4. Proceed to the next tab to create your template
            """)
            st.markdown('</div>', unsafe_allow_html=True)

            if st.session_state.rfp_text:
                if st.button("Analyze RFP", type="primary"):
                    with st.spinner("Analyzing RFP..."):
                        if st.session_state.generator:
                            # Pass cleaned RFP text to analyze_rfp
                            rfp_analysis = st.session_state.generator.analyze_rfp(remove_problematic_chars(st.session_state.rfp_text))
                            st.session_state.rfp_analysis = rfp_analysis # Store cleaned analysis

                            # Extract required sections (uses cleaned analysis internally)
                            required_sections = st.session_state.generator.extract_required_sections(rfp_analysis)
                            st.session_state.required_sections = required_sections

                            # Extract mandatory criteria (uses cleaned analysis internally)
                            mandatory_criteria = st.session_state.generator.extract_mandatory_criteria(rfp_analysis)
                            st.session_state.mandatory_criteria = mandatory_criteria

                            # Extract deadlines (uses cleaned analysis internally)
                            deadlines = st.session_state.generator.extract_deadlines(rfp_analysis)
                            st.session_state.deadlines = deadlines

                            # Extract deliverables (uses cleaned analysis internally)
                            deliverables = st.session_state.generator.extract_deliverables(rfp_analysis)
                            st.session_state.deliverables = deliverables

                            # Assess compliance (uses cleaned analysis internally)
                            internal_capabilities = st.session_state.config.get("internal_capabilities", {})
                            compliance_assessment = st.session_state.generator.assess_compliance(rfp_analysis, internal_capabilities)
                            st.session_state.compliance_assessment = compliance_assessment

                            st.success("RFP Analysis Complete")
                            st.markdown("### Key Insights")
                            st.markdown("#### Mandatory Criteria")
                            if st.session_state.mandatory_criteria:
                                # Display cleaned mandatory criteria
                                st.markdown("\n".join([f"- {remove_problematic_chars(item)}" for item in st.session_state.mandatory_criteria]))
                            else:
                                st.markdown("No mandatory criteria found.")

                            st.markdown("#### Deadlines")
                            if st.session_state.deadlines:
                                # Display cleaned deadlines
                                st.markdown("\n".join([f"- {remove_problematic_chars(item)}" for item in st.session_state.deadlines]))
                            else:
                                st.markdown("No deadlines found.")

                            st.markdown("#### Deliverables")
                            if st.session_state.deliverables:
                                # Display cleaned deliverables
                                st.markdown("\n".join([f"- {remove_problematic_chars(item)}" for item in st.session_state.deliverables]))
                            else:
                                st.markdown("No deliverables found.")

                            st.markdown("#### Compliance Assessment")
                            # Display cleaned compliance assessment
                            st.markdown(remove_problematic_chars(st.session_state.compliance_assessment))

                            st.markdown("#### Full RFP Analysis")
                            # Display cleaned full analysis
                            st.write(remove_problematic_chars(rfp_analysis))
                        else:
                            st.error("OpenAI API key is not configured")

    # Tab 2: Template Creation
    with tabs[1]:
        st.header("Create Proposal Template")

        if not st.session_state.rfp_analysis:
            st.warning("Please upload and analyze an RFP first.")
        else:
            col1, col2 = st.columns([2, 1])

            with col1:
                st.markdown("### Define Template Sections")
                st.markdown("Select sections from the suggestions below or add your own.")

                st.markdown("#### Sections from Current RFP Analysis")
                if hasattr(st.session_state, 'required_sections') and st.session_state.required_sections:
                    with st.expander("Select sections identified in this RFP", expanded=True):
                        for section in st.session_state.required_sections:
                            # Ensure section name is cleaned for display and key
                            cleaned_section = remove_problematic_chars(section)
                            already_added = cleaned_section in st.session_state.template_sections
                            # Use a consistent key structure with cleaned name
                            is_selected = st.checkbox(cleaned_section, value=already_added, key=f"rfp_section_select_{cleaned_section}")
                            if is_selected and cleaned_section not in st.session_state.template_sections:
                                st.session_state.template_sections.append(cleaned_section)
                                # st.rerun() # Rerun to update the list visually
                            elif not is_selected and cleaned_section in st.session_state.template_sections:
                                st.session_state.template_sections.remove(cleaned_section)
                                # st.rerun() # Rerun to update the list visually
                else:
                    st.caption("No specific sections were automatically extracted from the RFP analysis.")

                st.markdown("---")

                st.markdown("#### Add Custom Section")
                new_section_name = st.text_input("Enter custom section name:", key="new_section_name_input")

                if st.button("Add Custom Section", type="secondary", key="add_custom_section_button"):
                    if new_section_name:
                        # Clean the new section name before adding
                        normalized_new_section = remove_problematic_chars(new_section_name.strip().title())
                        if normalized_new_section and normalized_new_section not in st.session_state.template_sections:
                            st.session_state.template_sections.append(normalized_new_section)
                            st.success(f"Section '{normalized_new_section}' added to template.")
                            st.session_state.new_section_name_input = "" # Clear input field
                            st.rerun() # Rerun to update the list immediately
                        elif normalized_new_section in st.session_state.template_sections:
                            st.warning(f"Section '{normalized_new_section}' already exists in template.")
                        else:
                             st.warning("Please provide a valid section name.")
                    else:
                        st.warning("Please provide a section name.")

                st.markdown("---")
                st.markdown("#### Current Proposal Template Sections")
                # Use a unique key for each item in the list for removal
                # Iterate over a copy of the list if modifying it during iteration
                current_sections = st.session_state.template_sections[:]
                for i, section in enumerate(current_sections):
                    sec_col1, sec_col2 = st.columns([4, 1])
                    with sec_col1:
                        # Display cleaned section name
                        st.write(f"{i+1}. {remove_problematic_chars(section)}")
                    with sec_col2:
                         # Use a unique key for each button based on index and cleaned section name
                         if st.button("Remove", key=f"remove_template_section_{i}_{remove_problematic_chars(section)}"):
                             st.session_state.template_sections.pop(i)
                             st.rerun() # Rerun to update the list
                             break # Exit loop after removing to avoid index issues


                if not st.session_state.template_sections:
                    st.info("No sections selected for the template yet.")

                st.markdown("---")
                if st.session_state.template_sections:
                    if st.button("Confirm Sections & Proceed to Generate", type="primary", key="confirm_template_button"):
                        st.session_state.template_created = True
                        st.success("Template sections confirmed. Proceed to the 'Generate Proposal' tab.")
                        # st.rerun() # Rerun to move to the next tab implicitly or explicitly
                else:
                    st.button("Confirm Sections & Proceed to Generate", type="primary", key="confirm_template_button_disabled", disabled=True)

            with col2:
                st.markdown('<div class="info-box sidebar-content">', unsafe_allow_html=True)
                st.markdown("### 📝 Template Instructions")
                st.markdown("""
                1.  **Select Sections:** Choose from:
                    * Sections identified in the current RFP.
                2.  **Add Custom:** Input any additional sections needed.
                3.  **Review:** Check the 'Current Proposal Template Sections' list.
                4.  **Remove:** Use the 'Remove' button next to any section you don't want.
                5.  **Confirm:** Click 'Confirm Sections & Proceed' when ready.
                """)
                st.markdown('</div>', unsafe_allow_html=True)


    # Tab 3: Generate Proposal
    with tabs[2]:
        st.header("Generate Proposal")

        if not st.session_state.template_created:
            st.warning("Please create a template first.")
        elif not st.session_state.generator:
             st.warning("OpenAI API key is not configured. Please configure it to generate proposals.")
        else:
            col1, col2 = st.columns([1, 1])

            with col1:
                st.markdown("### Proposal Configuration")

                # Use the client_name from session_state if available, otherwise default
                # Clean client name input
                client_name_gen = remove_problematic_chars(st.text_input("Client Name", st.session_state.get('client_name_input', "Client Organization"), key="client_name_input_gen"))
                st.session_state.client_name_input = client_name_gen # Keep consistent client name across tabs


                company_info = {
                    "name": st.session_state.config["company_info"]["name"], # Assuming company name is safe ASCII
                    # Clean differentiators input
                    "differentiators": remove_problematic_chars(st.text_area("Company Differentiators",
                                                     st.session_state.get('differentiators_input', "Enter key differentiators that make your company stand out"), # Preserve input
                                                     key="differentiators_input"))
                }

                if st.button("Generate Proposal", type="primary"):
                    with st.spinner("Generating proposal..."):
                        try:
                            # Generate proposal using template sections
                            proposal_data = st.session_state.generator.generate_full_proposal(
                                st.session_state.rfp_text, # rfp_text is already cleaned on upload
                                client_name_gen, # Use the cleaned client name from input
                                company_info, # Company info contains cleaned differentiators
                                st.session_state.template_sections # Template sections are already cleaned
                            )
                            st.session_state.proposal_data = proposal_data
                            st.success("Proposal generated successfully!")
                            st.rerun() # Rerun to display the preview tabs
                        except Exception as e:
                            st.error(f"Error generating proposal: {str(e)}")

            with col2:
                st.markdown("### Generation Controls")
                st.markdown("""
                The proposal generation process:
                1. Uses the RFP analysis to identify required sections
                2. Embeds the RFP analysis for similarity search
                3. Retrieves relevant content from the knowledge base
                4. Generates sections that strictly adhere to RFP requirements
                """)

            if st.session_state.proposal_data and st.session_state.proposal_data["sections"]:
                st.markdown("---")
                st.header("Proposal Preview")

                section_names = list(st.session_state.proposal_data["sections"].keys())

                # Use a list of section names to create tabs
                section_tabs = st.tabs(section_names)

                for i, section_name in enumerate(section_names):
                    content = st.session_state.proposal_data["sections"][section_name]
                    with section_tabs[i]:
                        # Display cleaned content in the preview
                        st.markdown(remove_problematic_chars(content))

                        st.markdown("---")
                        col1, col2 = st.columns([3, 1])

                        with col1:
                            feedback = st.text_area(
                                "Provide feedback to improve this section:",
                                key=f"feedback_{remove_problematic_chars(section_name)}" # Unique key with cleaned name
                            )

                        with col2:
                            st.markdown("<br>", unsafe_allow_html=True) # Add some space above the button
                            if st.button("Update Section", key=f"update_{remove_problematic_chars(section_name)}"):
                                if feedback:
                                    try:
                                        with st.spinner(f"Updating '{remove_problematic_chars(section_name)}' based on feedback..."):
                                            if not st.session_state.generator:
                                                raise Exception("OpenAI API key is not configured")

                                            # Pass cleaned inputs to refine_section
                                            refined_content = st.session_state.generator.refine_section(
                                                section_name, # section_name is already cleaned in proposal_data
                                                content, # content is already cleaned in proposal_data
                                                feedback, # feedback is cleaned within refine_section
                                                st.session_state.proposal_data.get('client_name', 'Client') # client_name is cleaned in proposal_data
                                            )

                                            # Store cleaned refined content
                                            st.session_state.proposal_data["sections"][section_name] = remove_problematic_chars(refined_content)

                                            st.rerun() # Rerun to show the updated section
                                    except Exception as e:
                                        st.error(f"Error updating section: {str(e)}")
                                else:
                                    st.warning("Please provide feedback to update the section.")

            else:
                st.info("No sections available. Generate your proposal first.")


    # Tab 4: Export
    with tabs[3]:
        st.header("Export Your Proposal")

        if not st.session_state.proposal_data or not st.session_state.proposal_data["sections"]:
            st.warning("Please generate your proposal first.")
        elif not st.session_state.generator:
             st.warning("OpenAI API key is not configured. Please configure it to generate proposals.")
        else:
            col1, col2 = st.columns([2, 1])

            with col1:
                st.markdown("### Export Settings")

                company_name = st.session_state.config["company_info"]["name"]
                # Use the client name from the generated proposal data (already cleaned)
                client_name_export = st.session_state.proposal_data.get("client_name", "Client")

                uploaded_logo = st.file_uploader("Upload Company Logo (optional)", type=["png", "jpg", "jpeg"])
                logo_path = None
                if uploaded_logo:
                    image = Image.open(uploaded_logo)
                    logo_path = "temp_logo.png" # Save to a temporary file
                    try:
                        image.save(logo_path)
                    except Exception as e:
                        st.error(f"Error saving logo file: {e}")
                        logo_path = None # Reset logo_path if saving fails


                export_format = st.selectbox(
                    "Export Format",
                    ["Word Document (.docx)", "PDF Document (.pdf)", "Markdown (.md)"],
                    key="export_format_select"
                )

                if st.button("Export", type="primary"):
                    with st.spinner(f"Exporting proposal as {export_format}..."):
                        try:
                            # Ensure output directory exists
                            output_dir = "exported_proposals"
                            if not os.path.exists(output_dir):
                                os.makedirs(output_dir)

                            if export_format == "Word Document (.docx)":
                                # Ensure filename is safe
                                output_filename = f"Proposal_for_{remove_problematic_chars(client_name_export).replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
                                output_path = os.path.join(output_dir, output_filename)
                                final_path = export_to_word(
                                    st.session_state.proposal_data, # Data contains cleaned text
                                    company_name, # company_name is assumed safe or cleaned in export_to_word
                                    client_name_export, # client_name_export is already cleaned
                                    output_path,
                                    logo_path
                                )
                                if final_path and os.path.exists(final_path):
                                    with open(final_path, "rb") as file:
                                        st.download_button(
                                            label="Download Word Document",
                                            data=file,
                                            file_name=output_filename,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                        )
                                else:
                                     st.error("Failed to create Word document.")


                            elif export_format == "PDF Document (.pdf)":
                                # Ensure filename is safe
                                output_filename = f"Proposal_for_{remove_problematic_chars(client_name_export).replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
                                output_path = os.path.join(output_dir, output_filename)
                                # Ensure client_name_export is passed as a string (already cleaned)
                                final_path = export_to_pdf(
                                    st.session_state.proposal_data, # Data contains cleaned text
                                    company_name, # company_name is assumed safe or cleaned in export_to_pdf
                                    client_name_export, # client_name_export is already cleaned
                                    output_path,
                                    logo_path
                                )
                                if final_path and os.path.exists(final_path):
                                    with open(final_path, "rb") as file:
                                        st.download_button(
                                            label="Download PDF",
                                            data=file,
                                            file_name=output_filename,
                                            mime="application/pdf"
                                        )
                                else:
                                     st.error("Failed to create PDF document. Ensure 'fpdf' is installed.")


                            else: # Markdown format
                                # Ensure filename is safe
                                output_filename = "Proposal_for_{}_{}.md".format(
                                    remove_problematic_chars(client_name_export).replace(' ', '_') if client_name_export else 'Client',
                                    datetime.now().strftime('%Y%m%d%H%M%S')
                                )
                                md_content = f"# Proposal for {remove_problematic_chars(client_name_export)}\n\n" # Clean client name
                                for section_name, content in st.session_state.proposal_data["sections"].items():
                                    # Clean section name and content for markdown
                                    md_content += f"## {remove_problematic_chars(section_name)}\n\n{remove_problematic_chars(content)}\n\n"

                                st.download_button(
                                    label="Download Markdown File",
                                    data=md_content,
                                    file_name=output_filename,
                                    mime="text/markdown"
                                )

                            st.success("Export completed successfully!")

                        except Exception as e:
                            st.error(f"An unexpected error occurred during export: {str(e)}")
                        finally:
                            # Clean up the temporary logo file after export attempt
                            if logo_path and os.path.exists(logo_path):
                                os.remove(logo_path)


            with col2:
                st.markdown("### Export Options")
                st.markdown("""
                You can export your proposal in multiple formats:
                1. **Word Document**: Professional document with formatting
                2. **PDF Document**: Fixed-format document for printing or sharing
                3. **Markdown**: Text-based format for easy editing
                """)

    # Tab 5: Advanced Analysis
    with tabs[4]:
        st.header("Advanced Proposal Analysis")

        if not st.session_state.proposal_data or not st.session_state.proposal_data["sections"]:
            st.warning("Please generate your proposal first.")
        elif not st.session_state.generator:
             st.warning("OpenAI API key is not configured. Please configure it to generate analysis.")
        else:
            if st.button("Generate Advanced Analysis", type="primary", key="advanced_analysis_button"):
                with st.spinner("Generating advanced analysis..."):
                    try:
                        internal_capabilities = st.session_state.config.get("internal_capabilities", {})
                        # Pass cleaned data to generate_advanced_analysis
                        advanced_analysis = st.session_state.generator.generate_advanced_analysis(
                            st.session_state.proposal_data, # Data contains cleaned text
                            st.session_state.rfp_analysis, # rfp_analysis is already cleaned
                            internal_capabilities, # internal_capabilities are cleaned within generate_advanced_analysis
                            st.session_state.proposal_data.get('client_name', 'Client') # client_name is cleaned in proposal_data
                        )
                        st.session_state.advanced_analysis = advanced_analysis # Store cleaned analysis results
                        st.success("Advanced Analysis Complete")
                        st.rerun() # Rerun to display results
                    except Exception as e:
                        st.error(f"Error generating advanced analysis: {str(e)}")

            if st.session_state.advanced_analysis and any(st.session_state.advanced_analysis.values()):
                st.markdown("### Advanced Analysis Results")

                if st.session_state.advanced_analysis.get("compliance_matrix"):
                    st.markdown("#### Compliance Matrix")
                    # Display cleaned compliance matrix
                    st.markdown(remove_problematic_chars(st.session_state.advanced_analysis["compliance_matrix"]))

                if st.session_state.advanced_analysis.get("risk_assessment"):
                    st.markdown("#### Risk Assessment")
                    # Display cleaned risk assessment
                    st.markdown(remove_problematic_chars(st.session_state.advanced_analysis["risk_assessment"]))

                if st.session_state.advanced_analysis.get("alignment_assessment"):
                    st.markdown("#### Alignment Assessment")
                    # Display cleaned alignment assessment
                    st.markdown(remove_problematic_chars(st.session_state.advanced_analysis["alignment_assessment"]))

                if st.session_state.advanced_analysis.get("compliance_assessment"):
                    st.markdown("#### Compliance Assessment")
                    # Display cleaned compliance assessment
                    st.markdown(remove_problematic_chars(st.session_state.advanced_analysis["compliance_assessment"]))

                if st.session_state.advanced_analysis.get("quality_assurance"):
                    st.markdown("#### Quality Assurance")
                    # Display cleaned quality assurance
                    st.markdown(remove_problematic_chars(st.session_state.advanced_analysis["quality_assurance"]))
            elif st.session_state.get("advanced_analysis_button"): # Only show if button was clicked but no results
                 st.info("Click 'Generate Advanced Analysis' to see results.")


    # Tab 6: Vendor Proposal Evaluation
    with tabs[5]: # Adjust index if you added/removed tabs
        st.header("Vendor Proposal Evaluation")

        if not st.session_state.rfp_analysis:
            st.warning("Please upload and analyze an RFP first (Tab 1).")
        elif not st.session_state.generator:
             st.warning("OpenAI API key is not configured. Please configure it to evaluate vendor proposals.")
        else:
            # --- Dynamic Weightage Configuration UI ---
            st.markdown("---")
            st.subheader("⚙️ Configure Scoring Weightage")

            # Get default weights from config or use session state if already modified
            # Initialize session state for dynamic weights if not already present
            # Ensure scoring_system and weighting keys exist before accessing
            if 'dynamic_weights' not in st.session_state or not st.session_state.dynamic_weights:
                 # Use a copy to avoid modifying the original config in session state directly
                 st.session_state.dynamic_weights = st.session_state.config.get('scoring_system', {}).get('weighting', {}).copy()
                 if not st.session_state.dynamic_weights:
                      # Fallback if weighting is also missing
                      st.session_state.dynamic_weights = {
                         "requirement_match": 0.4,
                         "compliance": 0.25,
                         "quality": 0.2,
                         "alignment": 0.15,
                         "risk": 0.1
                      }
                      st.warning("Scoring system weights not found in config. Using default weights.")


            # Ensure consistent number of columns based on the number of metrics
            num_metrics = len(st.session_state.dynamic_weights)
            # Use a fixed number of columns (e.g., 4 or 5) and wrap metrics
            num_cols_for_weights = 4
            cols = st.columns(num_cols_for_weights)

            st.markdown("Enter weights (e.g., as decimals summing to 1.0 or percentages summing to 100):")

            total_weight_sum = 0.0
            metric_inputs = {} # Dictionary to hold the st.number_input widgets

            metrics_list = list(st.session_state.dynamic_weights.keys()) # Get metrics from dynamic_weights
            for i, metric in enumerate(metrics_list):
                # Calculate the column index
                col_index = i % num_cols_for_weights
                with cols[col_index]:
                    # Use the value from session_state.dynamic_weights for initialization
                    current_weight = st.session_state.dynamic_weights.get(metric, 0.0) # Default to 0.0 if metric somehow missing
                    metric_inputs[metric] = st.number_input(
                        f"{metric.replace('_', ' ').title()}",
                        min_value=0.0,
                        # max_value=100.0, # Allow input up to 100 for percentage-like input - removed max for flexibility
                        value=current_weight,
                        step=0.01,
                        format="%.2f",
                        key=f"weight_input_{metric}" # Unique key for each input
                    )
                    # Update the session state with the new value immediately when input changes
                    st.session_state.dynamic_weights[metric] = metric_inputs[metric]
                    total_weight_sum += metric_inputs[metric]


            # Display the current sum of weights
            st.info(f"Current total weight sum: {total_weight_sum:.2f}")
            # Add a visual cue if the sum is not close to 1 or 100
            if abs(total_weight_sum - 1.0) > 0.01 and abs(total_weight_sum - 100.0) > 1:
                 st.warning("Weights do not sum to 1.0 or 100.0. The calculated score will be a sum of (score * weight).")


            st.markdown("---")
            # --- End Dynamic Weightage Configuration UI ---

            uploaded_vendor_proposal = st.file_uploader("Upload Vendor Proposal", type=["docx", "pdf", "txt", "md"], key="vendor_proposal_upload")

            if uploaded_vendor_proposal:
                # ... (file processing logic as before) ...
                try:
                    # Use a unique key for vendor proposal text state
                    # Check if a file is uploaded AND it's a new file or the content needs reprocessing
                    if ('vendor_proposal_text' not in st.session_state or
                        st.session_state.get('processed_vendor_file') != uploaded_vendor_proposal.name or
                        st.session_state.get('processed_vendor_file_size') != uploaded_vendor_proposal.size): # Check file size too

                         temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_vendor_proposal.name.split('.')[-1]}")
                         temp_file.write(uploaded_vendor_proposal.getvalue())
                         temp_file.close()
                         proposal_text = process_rfp(temp_file.name)
                         st.session_state.vendor_proposal_text = proposal_text # Store cleaned text
                         st.session_state.processed_vendor_file = uploaded_vendor_proposal.name # Track processed file name
                         st.session_state.processed_vendor_file_size = uploaded_vendor_proposal.size # Track processed file size
                         st.session_state.vendor_analysis = None # Reset analysis when new file uploaded
                         st.session_state.vendor_score_results = None # Reset scores
                         st.session_state.vendor_gaps_risks = None # Reset gaps/risks
                         st.success(f"Successfully processed {uploaded_vendor_proposal.name}")
                         if os.path.exists(temp_file.name):
                             os.unlink(temp_file.name)
                    elif st.session_state.get('vendor_proposal_text'):
                         # If the file is the same and text is already loaded, just show the preview
                         with st.expander("Preview Vendor Proposal Content", expanded=False):
                             # Display cleaned text in the preview
                             st.text_area("Proposal Text", remove_problematic_chars(st.session_state.vendor_proposal_text), height=300, key="vendor_proposal_preview")


                except Exception as e:
                    st.error(f"Error processing file: {str(e)}")
                    # Clean up temp file if it was created during an error
                    if 'temp_file' in locals() and os.path.exists(temp_file.name):
                        os.unlink(temp_file.name)

            # Only show analysis button if vendor text is loaded
            if st.session_state.get('vendor_proposal_text'):
                # Use the client name from the generated proposal data if available, otherwise default
                # Clean client name input
                client_name_eval = remove_problematic_chars(st.text_input("Client Name (for analysis context)", st.session_state.proposal_data.get('client_name', "Client Organization"), key="client_name_eval"))

                if st.button("Analyze Vendor Proposal", type="primary", key="analyze_vendor_button"):
                    with st.spinner("Analyzing vendor proposal... This may take a moment."):
                        try:
                            # Call analyze_vendor_proposal.
                            # Pass the *current* scoring system configuration to the LLM for scoring guidance.
                            # The weights themselves are primarily used for the calculation *after* the LLM scores.
                            current_scoring_config = {
                                "weighting": st.session_state.dynamic_weights,
                                "grading_scale": st.session_state.config.get('scoring_system', {}).get('grading_scale', {}) # Safely access grading_scale
                            }
                            # Pass cleaned inputs to analyze_vendor_proposal
                            analysis = st.session_state.generator.analyze_vendor_proposal(
                                remove_problematic_chars(st.session_state.vendor_proposal_text),
                                remove_problematic_chars(st.session_state.rfp_analysis),
                                client_name_eval, # client_name_eval is already cleaned
                                current_scoring_config # Pass the current config to the LLM
                            )
                            st.session_state.vendor_analysis = analysis # Store cleaned raw analysis

                            # Use the dynamic weights from session_state for calculation
                            # Pass cleaned analysis text to calculate_weighted_score
                            weighted_score, individual_scores, grade = st.session_state.generator.calculate_weighted_score(
                                analysis, # analysis is already cleaned
                                {"weighting": st.session_state.dynamic_weights, "grading_scale": st.session_state.config.get('scoring_system', {}).get('grading_scale', {})} # Safely access grading_scale
                            )
                            st.session_state.vendor_score_results = {
                                "weighted_score": weighted_score,
                                "individual_scores": individual_scores,
                                "grade": grade
                            }

                            # Identify gaps and risks (Optional but useful)
                            # Pass cleaned inputs to identify_gaps_and_risks
                            rfp_requirements_text = remove_problematic_chars(st.session_state.rfp_analysis) # Use cleaned full analysis for context
                            gaps, risks = st.session_state.generator.identify_gaps_and_risks(
                                remove_problematic_chars(st.session_state.vendor_proposal_text),
                                rfp_requirements_text
                            )
                            st.session_state.vendor_gaps_risks = {"gaps": gaps, "risks": risks} # Gaps/risks are already cleaned

                            st.success("Analysis and Scoring Complete!")
                            st.rerun() # Rerun to display results cleanly below

                        except Exception as e:
                            st.error(f"Error analyzing proposal: {str(e)}")
                            st.session_state.vendor_analysis = remove_problematic_chars(f"Analysis Error: {str(e)}") # Store cleaned error message
                            st.session_state.vendor_score_results = None
                            st.session_state.vendor_gaps_risks = None

            # Display analysis results, scores, and gaps/risks if available
            if st.session_state.get('vendor_analysis'):
                st.markdown("---")
                st.header("Vendor Analysis Results")

                # Display Scoring Summary First
                if st.session_state.get('vendor_score_results'):
                    score_results = st.session_state.vendor_score_results
                    st.subheader("📊 Scoring Summary (Based on Configured Weights)")
                    # Use the calculated weighted score from the results
                    if score_results['weighted_score'] is not None:
                         # We'll display the calculated score after normalization in calculate_weighted_score.
                         st.metric(label="Overall Weighted Score (Normalized)", value=f"{score_results['weighted_score']:.2f}")
                         # Display cleaned grade
                         st.metric(label="Grade", value=remove_problematic_chars(score_results['grade'] or "N/A"))

                         st.markdown("##### Individual Metric Scores (as scored by AI - 0-100):")
                         # Ensure individual_scores is not None before iterating
                         if score_results.get('individual_scores'):
                              # Sort metrics alphabetically for consistent display
                              metrics_to_display = sorted(score_results['individual_scores'].keys())
                              cols = st.columns(min(len(metrics_to_display), 5)) # Limit columns for display
                              i = 0
                              for metric in metrics_to_display:
                                  score = score_results['individual_scores'].get(metric) # Use .get for safety
                                  with cols[i % len(cols)]:
                                         # Display cleaned metric name and score
                                         st.metric(label=remove_problematic_chars(metric.replace('_', ' ').title()), value=str(score) if score is not None else "N/A")
                                  i += 1
                              st.caption("Note: Individual scores are AI's raw assessment (0-100); Overall Score is calculated using the configured weights above.")
                         else:
                              st.info("No individual metric scores were extracted from the AI analysis.")

                    else:
                         st.warning("Could not calculate weighted score. Scores may be missing in the analysis or weights not configured.")
                         # Display cleaned individual scores
                         st.write("Individual Scores Found:", remove_problematic_chars(str(score_results.get('individual_scores', "N/A"))))

                # Display Gaps and Risks
                if st.session_state.get('vendor_gaps_risks'):
                     gaps_risks = st.session_state.vendor_gaps_risks
                     # Check if either gaps or risks list is non-empty
                     if gaps_risks.get('gaps') or gaps_risks.get('risks'):
                          st.subheader("⚠️ Identified Gaps & Risks (Beta)")
                          if gaps_risks.get('gaps'):
                               st.markdown("##### Gaps:")
                               if gaps_risks['gaps']:
                                    for gap in gaps_risks['gaps']:
                                         # Display cleaned gaps
                                         st.markdown(f"- {remove_problematic_chars(gap)}")
                               else:
                                    st.info("No significant gaps automatically identified.")
                          if gaps_risks.get('risks'):
                               st.markdown("##### Risks:")
                               if gaps_risks['risks']:
                                    for risk in gaps_risks['risks']:
                                         # Display cleaned risks
                                         st.markdown(f"- {remove_problematic_chars(risk)}")
                               else:
                                     st.info("No significant risks automatically identified.")
                     # If both lists are empty, show a general message
                     elif gaps_risks.get('gaps') is not None and gaps_risks.get('risks') is not None:
                          st.info("No significant gaps or risks automatically identified based on similarity and keywords.")


                # Display the full AI analysis text
                st.subheader("🤖 Full AI Analysis Text")
                # Display cleaned vendor analysis
                st.markdown(remove_problematic_chars(st.session_state.vendor_analysis))

    # Tab 7: RFP Template Creator
    with tabs[6]:
        st.header("RFP Template Creator")

        col1, col2 = st.columns([2, 1])

        with col1:
            st.markdown("### Create RFP Template")
            # Clean company objectives input
            company_objectives = remove_problematic_chars(st.text_area("Company Objectives", height=200, key="objectives_input"))
            # Clean template type input
            template_type = remove_problematic_chars(st.selectbox("Template Type", ["Standard RFP", "Technical RFP", "Commercial RFP", "Custom"], key="template_type_select"))

            if template_type == "Custom":
                # Clean custom template name input
                custom_template_name = remove_problematic_chars(st.text_input("Custom Template Name", key="custom_template_name"))
            else:
                custom_template_name = ""

            if st.button("Generate RFP Template", type="primary", key="generate_template_button"):
                with st.spinner("Generating RFP template..."):
                    try:
                        # Check if OpenAI API key is configured
                        openai_key = st.session_state.config["api_keys"]["openai_key"]
                        if not openai_key:
                             st.error("OpenAI API key is not configured. Please add it to config.json or set the OPENAI_API_KEY environment variable.")
                             st.stop() # Stop execution if key is missing

                        drafter = SpecialistRAGDrafter(openai_key)
                        # Pass cleaned inputs to generate_rfp_template
                        template_content = drafter.generate_rfp_template(company_objectives, template_type if template_type != "Custom" else custom_template_name)
                        st.session_state.rfp_template_content = template_content # Store cleaned template content
                        st.success("RFP Template generated successfully!")
                        st.rerun() # Rerun to display the preview
                    except Exception as e:
                        st.error(f"Error generating RFP template: {str(e)}")

        with col2:
            st.markdown('<div class="info-box">', unsafe_allow_html=True)
            st.markdown("### 📝 Template Creator Instructions")
            st.markdown("""
            1. Describe your company objectives in detail
            2. Select a template type or create a custom one
            3. Click "Generate RFP Template"
            4. Review and download the generated template
            """)
            st.markdown('</div>', unsafe_allow_html=True)

        if st.session_state.rfp_template_content:
            st.markdown("---")
            st.header("RFP Template Preview")
            # Display cleaned template content
            st.markdown(remove_problematic_chars(st.session_state.rfp_template_content))

            # Add download button for the template
            st.markdown("---")
            st.header("Download RFP Template")
            # Use the custom name if provided, otherwise the template type (already cleaned)
            template_filename_base = custom_template_name.replace(' ', '_') if template_type == 'Custom' and custom_template_name else template_type.replace(' ', '_')
            # Ensure filename is safe
            template_name = f"{remove_problematic_chars(template_filename_base)}_{datetime.now().strftime('%Y%m%d%H%M%S')}.md"
            st.download_button(
                label="Download Template",
                data=st.session_state.rfp_template_content, # Data is already cleaned
                file_name=template_name,
                mime="text/markdown",
                key="download_template_button"
            )


if __name__ == "__main__":
    main()
